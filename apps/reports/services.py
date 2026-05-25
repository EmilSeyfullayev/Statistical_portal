from pathlib import Path

from django.conf import settings
from django.utils import timezone
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

from apps.audit.models import AuditLog
from apps.audit.services import log_action
from apps.imports.models import TransitRecord
from apps.reports.models import ReportDownload


def generate_transit_docx(report):
    output_dir = Path(settings.REPORT_OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{report.slug}-{timezone.now():%Y%m%d-%H%M%S}.docx"
    output_path = output_dir / filename

    document = Document()
    document.add_heading(report.name, level=1)
    document.add_paragraph(f"Module: {report.submodule.module.name}")
    document.add_paragraph(f"Submodule: {report.submodule.name}")
    document.add_paragraph(f"Generated: {timezone.localtime():%Y-%m-%d %H:%M}")

    records = TransitRecord.objects.filter(source_file__submodule=report.submodule)
    document.add_heading("Summary", level=2)
    document.add_paragraph(f"Imported rows: {records.count()}")

    table = document.add_table(rows=1, cols=6)
    headers = ["Date", "Country", "Corridor", "Post", "Cargo", "Weight tons"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for record in records[:100]:
        cells = table.add_row().cells
        cells[0].text = record.date.isoformat() if record.date else ""
        cells[1].text = record.country
        cells[2].text = record.corridor
        cells[3].text = record.post
        cells[4].text = record.cargo_name
        cells[5].text = str(record.weight_tons or "")

    document.save(output_path)
    return output_path


def generate_transit_pdf(report):
    output_dir = Path(settings.REPORT_OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{report.slug}-{timezone.now():%Y%m%d-%H%M%S}.pdf"
    output_path = output_dir / filename

    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(str(output_path), pagesize=A4)
    records = TransitRecord.objects.filter(source_file__submodule=report.submodule)
    content = [
        Paragraph(report.name, styles["Title"]),
        Paragraph(f"Module: {report.submodule.module.name}", styles["Normal"]),
        Paragraph(f"Submodule: {report.submodule.name}", styles["Normal"]),
        Paragraph(f"Generated: {timezone.localtime():%Y-%m-%d %H:%M}", styles["Normal"]),
        Spacer(1, 12),
        Paragraph(f"Imported rows: {records.count()}", styles["Heading2"]),
    ]
    data = [["Date", "Country", "Corridor", "Post", "Cargo", "Weight tons"]]
    for record in records[:100]:
        data.append(
            [
                record.date.isoformat() if record.date else "",
                record.country,
                record.corridor,
                record.post,
                record.cargo_name,
                str(record.weight_tons or ""),
            ]
        )
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#174a7c")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )
    content.append(table)
    document.build(content)
    return output_path


REPORT_GENERATORS = {
    "transit_docx_v1": generate_transit_docx,
    "transit_pdf_v1": generate_transit_pdf,
}


def generate_report_file(report):
    try:
        generator = REPORT_GENERATORS[report.generator_key]
    except KeyError as exc:
        raise ValueError(f"No report generator registered for '{report.generator_key}'.") from exc
    return generator(report)


def record_download(request, report, generated_file_path):
    download = ReportDownload.objects.create(
        user=request.user if request.user.is_authenticated else None,
        report=report,
        module=report.submodule.module,
        submodule=report.submodule,
        format=report.format,
        generated_file_path=str(generated_file_path),
        request_metadata={
            "ip": request.META.get("REMOTE_ADDR", ""),
            "user_agent": request.META.get("HTTP_USER_AGENT", "")[:300],
        },
    )
    log_action(
        user=request.user,
        action_type="report_download",
        status=AuditLog.STATUS_SUCCESS,
        module=report.submodule.module,
        submodule=report.submodule,
        file_or_report=report.name,
        metadata={"download_id": download.id, "path": str(generated_file_path)},
    )
    return download
