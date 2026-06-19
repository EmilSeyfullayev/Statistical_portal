from copy import deepcopy
from datetime import date
from io import BytesIO

from django.db import connection
from django.http import QueryDict
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from apps.dashboard.transit_periods import (
    TransitDataPeriods,
    compute_transit_data_periods,
    format_dynamics_report_number,
    load_transit_data_periods,
)
from apps.dashboard.transit_value_translations import translate_transit_value
from apps.imports.services.goods_nomenclature_short import goods_nomenclature_short_data
from apps.imports.services.goods_nomenclature_short_English import goods_nomenclature_short_english_data
from apps.imports.services.goods_nomenclature_short_Russian import goods_nomenclature_short_russian_data

PAGE_SIZE = 100
DOWNLOAD_LIMIT = 500

TRANSIT_NATIVE_COLUMNS = [
    "goods_id",
    "giris_tarixi",
    "giris_nv_novu",
    "giris_go",
    "giris_dehliz",
    "bosaltma_yukleme_go",
    "cixis_tarixi",
    "cixis_nv_novu",
    "cixis_go",
    "cixis_dehliz",
    "mal_gonderen_olke",
    "mal_teyinat_olke",
    "mal_kodu",
    "mal_ceki_ton",
]

TRANSIT_FILTERS = {
    "date_from": "Giriş tarixi",
    "date_to": "Çıxış tarixi",
    "mal_gonderen_olke": "Mal göndərən ölkə",
    "mal_teyinat_olke": "Mal təyinat ölkə",
    "giris_dehliz": "Giriş dəhliz",
    "cixis_dehliz": "Çıxış dəhliz",
    "mal_kodu": "Mal kodu",
}

TRANSIT_FILTER_FIELDS = [
    {"name": "date_from", "label": "Giriş tarixi", "type": "date"},
    {"name": "date_to", "label": "Çıxış tarixi", "type": "date"},
    {"name": "mal_gonderen_olke", "label": "Mal göndərən ölkə", "type": "select"},
    {"name": "mal_teyinat_olke", "label": "Mal təyinat ölkə", "type": "select"},
    {"name": "giris_dehliz", "label": "Giriş dəhliz", "type": "select"},
    {"name": "cixis_dehliz", "label": "Çıxış dəhliz", "type": "select"},
    {"name": "mal_kodu", "label": "Mal kodu", "type": "select"},
]

TRANSIT_COLUMN_LABELS = {
    "goods_id": "Goods id",
    "giris_tarixi": "Giriş tarixi",
    "giris_nv_novu": "Giriş nv növü",
    "giris_go": "Giriş gö",
    "giris_dehliz": "Giriş dəhliz",
    "bosaltma_yukleme_go": "Boşaltma yükləmə gö",
    "cixis_tarixi": "Çıxış tarixi",
    "cixis_nv_novu": "Çıxış nv növü",
    "cixis_go": "Çıxış gö",
    "cixis_dehliz": "Çıxış dəhliz",
    "mal_gonderen_olke": "Mal göndərən ölkə",
    "mal_teyinat_olke": "Mal təyinat ölkə",
    "mal_kodu": "Mal kodu",
    "mal_ceki_ton": "Mal çəkisi ton",
}

DATE_COLUMNS = {"giris_tarixi", "cixis_tarixi"}

PORTAL_TABLE = "transit_data_for_portal"
PORTAL_SCHEMA = "transit"
DASHBOARD_TABLE = "transit_for_dashboard_on_ministry_portal"

PORTAL_COLUMNS = [
    "Giriş tarixi",
    "Giriş nəqliyyat növü",
    "Gömrük giriş postu",
    "Giriş nöqtəsi",
    "Boşaltma-yükləmə gömrük postu",
    "Çıxış tarixi",
    "Çıxış nəqliyyat növü",
    "Gömrük çıxış postu",
    "Çıxış nöqtəsi",
    "Göndərən ölkə",
    "Təyinat ölkə",
    "MAL_KODU",
    "Dəhliz (istiqamətlə)",
    "Dəhliz",
    "Giriş-Çıxış nəqliyyat növü",
    "Nəqliyyat növü",
    "Məhsulun adı",
    "Məhsul qrupu",
    "Məhsul adı (qısaldılmış)",
    "Məhsul qrupu (qısaldılmış)",
    "Başlangıc-Təyinat ölkəsi",
    "Yük həcmi (ton)",
]

PORTAL_FILTER_COLUMNS = {
    "date_from": "Giriş tarixi",
    "date_to": "Çıxış tarixi",
    "mal_gonderen_olke": "Göndərən ölkə",
    "mal_teyinat_olke": "Təyinat ölkə",
    "giris_dehliz": "Giriş nöqtəsi",
    "cixis_dehliz": "Çıxış nöqtəsi",
    "neqliyyat_novu": "Nəqliyyat növü",
    "mehsulun_adi": "Məhsul adı (qısaldılmış)",
    "mehsul_qrupu": "Məhsul qrupu (qısaldılmış)",
    "dehliz_istiqametle": "Dəhliz (istiqamətlə)",
    "dehliz": "Dəhliz",
}

PORTAL_FILTER_FIELDS = [
    {"name": "date_from", "label": "Giriş tarixi", "type": "date"},
    {"name": "date_to", "label": "Çıxış tarixi", "type": "date"},
    {"name": "mal_gonderen_olke", "label": "Göndərən ölkə", "type": "select"},
    {"name": "mal_teyinat_olke", "label": "Təyinat ölkə", "type": "select"},
    {"name": "giris_dehliz", "label": "Giriş nöqtəsi", "type": "select"},
    {"name": "cixis_dehliz", "label": "Çıxış nöqtəsi", "type": "select"},
    {"name": "dehliz_istiqametle", "label": "Dəhliz (istiqamətlə)", "type": "select"},
    {"name": "dehliz", "label": "Dəhliz", "type": "select"},
    {"name": "neqliyyat_novu", "label": "Nəqliyyat növü", "type": "select"},
    {"name": "mehsulun_adi", "label": "Məhsul adı (qısaldılmış)", "type": "select"},
    {"name": "mehsul_qrupu", "label": "Məhsul qrupu (qısaldılmış)", "type": "select"},
]

PORTAL_FILTERS = {field["name"]: field["label"] for field in PORTAL_FILTER_FIELDS}

PORTAL_DATE_COLUMNS = {"Giriş tarixi", "Çıxış tarixi"}
PORTAL_ORDER_COLUMN = "Çıxış tarixi"

DASHBOARD_COLUMNS = [
    "Gömrük giriş postu",
    "Çıxış tarixi",
    "Gömrük çıxış postu",
    "Göndərən ölkə",
    "Təyinat ölkə",
    "Dəhliz (istiqamətlə)",
    "Dəhliz",
    "Nəqliyyat növü",
    "Məhsul adı (qısaldılmış)",
    "Məhsul qrupu (qısaldılmış)",
    "Yük həcmi (ton)",
]

DASHBOARD_DEFAULT_TRANSPORTS = ["Avtomobil", "Dəmiryolu", "Hava"]
DASHBOARD_VALUE_COLUMN = "Yük həcmi (ton)"
DASHBOARD_DATE_COLUMN = "Çıxış tarixi"

FOREIGN_TRUCKS_SCHEMA = "foreign_trucks"
FOREIGN_TRUCKS_MERGED_TABLE = "foreign_trucks_merged"
FOREIGN_TRUCKS_COLUMNS = [
    "IDN",
    "CODE",
    "SHORT_NAME",
    "AVTO_NO",
    "ENTER_DATE",
    "CUST_NAME",
    "DATESIGN",
    "PERM_BLANK_NO",
    "PERMISSION_PRICE",
    "DIRECTION",
    "HES_NAME",
    "CONS_NAME",
    "WEIGHT",
    "TOTAL_WEIGHT",
    "FROMTO",
    "WIDTH",
    "HEIGHT",
    "WEIGHT_PER_AX",
    "PLACE_WHEEL_COUNT",
]
FOREIGN_TRUCKS_DATE_COLUMNS = {"ENTER_DATE", "DATESIGN"}
FOREIGN_TRUCKS_FILTER_COLUMNS = {
    "date_from": "ENTER_DATE",
    "date_to": "ENTER_DATE",
    "idn": "IDN",
    "code": "CODE",
    "cust_name": "CUST_NAME",
    "perm_blank_no": "PERM_BLANK_NO",
    "fromto": "FROMTO",
    "short_name": "SHORT_NAME",
    "avto_no": "AVTO_NO",
    "hes_name": "HES_NAME",
    "cons_name": "CONS_NAME",
}
FOREIGN_TRUCKS_FILTER_FIELDS = [
    {"name": "date_from", "label": "Daxil olma tarixi - başlanğıc", "type": "date"},
    {"name": "date_to", "label": "Daxil olma tarixi - son", "type": "date"},
    {"name": "idn", "label": "IDN", "type": "text"},
    {"name": "code", "label": "CODE", "type": "text"},
    {"name": "cust_name", "label": "Gömrük postu", "type": "select"},
    {"name": "perm_blank_no", "label": "İcazə blankın nömrəsi", "type": "text"},
    {"name": "fromto", "label": "Başlanğıc-təyinat ölkəsi", "type": "combobox"},
    {"name": "short_name", "label": "Daşıyıcı Mənsubiyyət ölkəsi", "type": "select"},
    {"name": "avto_no", "label": "Avtomobil nömrəsi", "type": "text"},
    {"name": "hes_name", "label": "İcazə növü", "type": "select"},
    {"name": "cons_name", "label": "Güzəşt növü", "type": "select"},
]
FOREIGN_TRUCKS_FILTERS = {field["name"]: field["label"] for field in FOREIGN_TRUCKS_FILTER_FIELDS}
FOREIGN_TRUCKS_SELECT_FILTERS = [
    field["name"] for field in FOREIGN_TRUCKS_FILTER_FIELDS if field["type"] == "select"
]
FOREIGN_TRUCKS_COMBOBOX_FILTERS = [
    field["name"] for field in FOREIGN_TRUCKS_FILTER_FIELDS if field["type"] == "combobox"
]
FOREIGN_TRUCKS_TEXT_FILTERS = [
    field["name"] for field in FOREIGN_TRUCKS_FILTER_FIELDS if field["type"] == "text"
]

DASHBOARD_FILTER_COLUMNS = {
    "date_from": DASHBOARD_DATE_COLUMN,
    "date_to": DASHBOARD_DATE_COLUMN,
    "neqliyyat_novu": "Nəqliyyat növü",
    "dehliz": "Dəhliz",
    "dehliz_istiqametle": "Dəhliz (istiqamətlə)",
    "mal_gonderen_olke": "Göndərən ölkə",
    "mal_teyinat_olke": "Təyinat ölkə",
    "mehsul_qrupu": "Məhsul qrupu (qısaldılmış)",
    "giris_go": "Gömrük giriş postu",
    "cixis_go": "Gömrük çıxış postu",
}

DASHBOARD_FILTER_FIELDS = [
    {"name": "date_from", "label": "Çıxış tarixi - başlanğıc", "type": "date"},
    {"name": "date_to", "label": "Çıxış tarixi - son", "type": "date"},
    {"name": "neqliyyat_novu", "label": "Nəqliyyat növü", "type": "multi_select"},
    {"name": "dehliz", "label": "Dəhliz", "type": "select"},
    {"name": "dehliz_istiqametle", "label": "Dəhliz (istiqamətlə)", "type": "select"},
    {"name": "mal_gonderen_olke", "label": "Göndərən ölkə", "type": "select"},
    {"name": "mal_teyinat_olke", "label": "Təyinat ölkə", "type": "select"},
    {"name": "mehsul_qrupu", "label": "Məhsul qrupu", "type": "select"},
    {"name": "giris_go", "label": "Gömrük giriş postu", "type": "select"},
    {"name": "cixis_go", "label": "Gömrük çıxış postu", "type": "select"},
]

REPORT_TOP_LIMIT = 10
REPORT_POST_LIMIT = 100


def quote_name(name):
    return connection.ops.quote_name(name)


def enrich_filter_fields(field_defs, filters, filter_options):
    return [
        {
            **field,
            "value": filters.get(field["name"], ""),
            "options": filter_options.get(field["name"], []),
        }
        for field in field_defs
    ]


def get_available_transit_columns():
    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = 'transit'
              AND table_name = 'transit_merged'
            """
        )
        available = {row[0] for row in cursor.fetchall()}
    return [column for column in TRANSIT_NATIVE_COLUMNS if column in available]


def transit_filter_options(columns):
    option_columns = [
        "mal_gonderen_olke",
        "mal_teyinat_olke",
        "giris_dehliz",
        "cixis_dehliz",
        "mal_kodu",
    ]
    options = {}
    with connection.cursor() as cursor:
        for column in option_columns:
            if column not in columns:
                options[column] = []
                continue
            quoted_column = quote_name(column)
            cursor.execute(
                f"""
                SELECT DISTINCT {quoted_column}
                FROM transit.transit_merged
                WHERE {quoted_column} IS NOT NULL AND btrim({quoted_column}) <> ''
                ORDER BY {quoted_column}
                LIMIT 500
                """
            )
            options[column] = [row[0] for row in cursor.fetchall()]
    return options


def build_transit_where(filters, columns):
    clauses = []
    params = []
    if filters.get("date_from") and "giris_tarixi" in columns:
        clauses.append(f"{quote_name('giris_tarixi')}::date >= %s")
        params.append(filters["date_from"])
    if filters.get("date_to"):
        if "merged_cixis_tarixi" in columns:
            clauses.append("merged_cixis_tarixi::date <= %s")
            params.append(filters["date_to"])
        elif "cixis_tarixi" in columns:
            clauses.append(f"{quote_name('cixis_tarixi')}::date <= %s")
            params.append(filters["date_to"])

    for column in [
        "mal_gonderen_olke",
        "mal_teyinat_olke",
        "giris_dehliz",
        "cixis_dehliz",
        "mal_kodu",
    ]:
        if column in columns and filters.get(column):
            clauses.append(f"{quote_name(column)} = %s")
            params.append(filters[column])
    where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    return where_sql, params


def clean_page_number(value):
    try:
        return max(1, int(value))
    except (TypeError, ValueError):
        return 1


def transit_select_sql(columns):
    select_parts = []
    for column in columns:
        quoted_column = quote_name(column)
        if column in DATE_COLUMNS:
            select_parts.append(f"NULLIF(left({quoted_column}::text, 10), '') AS {quoted_column}")
        else:
            select_parts.append(quoted_column)
    return ", ".join(select_parts)


def transit_column_labels(columns):
    return [TRANSIT_COLUMN_LABELS.get(column, column) for column in columns]


def pagination_items(page, total_pages):
    pages = {1, total_pages}
    pages.update(range(max(1, page - 2), min(total_pages, page + 2) + 1))
    items = []
    previous = 0
    for page_number in sorted(pages):
        if page_number - previous > 1:
            items.append({"ellipsis": True})
        items.append({"number": page_number, "current": page_number == page})
        previous = page_number
    return items


def transit_rows(params, *, limit=PAGE_SIZE, offset=0):
    columns = get_available_transit_columns()
    filters = {key: params.get(key, "").strip() for key in TRANSIT_FILTERS}
    where_sql, query_params = build_transit_where(filters, columns)
    select_columns = transit_select_sql(columns)
    order_sql = "ORDER BY merged_cixis_tarixi DESC NULLS LAST"

    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {select_columns}
            FROM transit.transit_merged
            {where_sql}
            {order_sql}
            LIMIT %s OFFSET %s
            """,
            [*query_params, limit, offset],
        )
        rows = [
            {"cells": row}
            for row in cursor.fetchall()
        ]
        cursor.execute(f"SELECT count(*) FROM transit.transit_merged {where_sql}", query_params)
        filtered_count = cursor.fetchone()[0]
    return columns, filters, rows, filtered_count


def get_transit_merged_context(params):
    page = clean_page_number(params.get("page"))
    offset = (page - 1) * PAGE_SIZE
    columns, filters, rows, filtered_count = transit_rows(params, limit=PAGE_SIZE, offset=offset)
    total_pages = max(1, (filtered_count + PAGE_SIZE - 1) // PAGE_SIZE)
    if page > total_pages:
        page = total_pages
        offset = (page - 1) * PAGE_SIZE
        columns, filters, rows, filtered_count = transit_rows(params, limit=PAGE_SIZE, offset=offset)
    base_query = QueryDict(mutable=True)
    for key, value in filters.items():
        if value:
            base_query[key] = value
    download_query = base_query.copy()
    download_query["download"] = "xlsx"
    filter_options = transit_filter_options(columns)

    return {
        "columns": columns,
        "column_labels": transit_column_labels(columns),
        "rows": rows,
        "filters": filters,
        "filter_labels": TRANSIT_FILTERS,
        "filter_options": filter_options,
        "filter_fields": enrich_filter_fields(TRANSIT_FILTER_FIELDS, filters, filter_options),
        "filtered_count": filtered_count,
        "limit": PAGE_SIZE,
        "page": page,
        "total_pages": total_pages,
        "has_previous": page > 1,
        "has_next": page < total_pages,
        "previous_page": page - 1,
        "next_page": page + 1,
        "pagination_items": pagination_items(page, total_pages),
        "page_start": offset + 1 if filtered_count else 0,
        "page_end": min(offset + len(rows), filtered_count),
        "query_string": base_query.urlencode(),
        "download_query_string": download_query.urlencode(),
        "download_limit": DOWNLOAD_LIMIT,
        "eyebrow": "Transit database",
        "heading": "Raw Transit Data",
    }


def get_transit_download_rows(params):
    columns = get_available_transit_columns()
    filters = {key: params.get(key, "").strip() for key in TRANSIT_FILTERS}
    where_sql, query_params = build_transit_where(filters, columns)
    select_columns = transit_select_sql(columns)
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {select_columns}
            FROM transit.transit_merged
            {where_sql}
            ORDER BY merged_cixis_tarixi DESC NULLS LAST
            LIMIT %s
            """,
            [*query_params, DOWNLOAD_LIMIT],
        )
        return transit_column_labels(columns), cursor.fetchall()


def get_available_portal_columns():
    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = %s
              AND table_name = %s
            """,
            [PORTAL_SCHEMA, PORTAL_TABLE],
        )
        available = {row[0] for row in cursor.fetchall()}
    return [column for column in PORTAL_COLUMNS if column in available]


def portal_filter_options(columns):
    option_keys = [
        field["name"]
        for field in PORTAL_FILTER_FIELDS
        if field["type"] == "select"
    ]
    options = {}
    qualified_table = f"{quote_name(PORTAL_SCHEMA)}.{quote_name(PORTAL_TABLE)}"
    with connection.cursor() as cursor:
        for key in option_keys:
            column = PORTAL_FILTER_COLUMNS[key]
            if column not in columns:
                options[key] = []
                continue
            quoted_column = quote_name(column)
            cursor.execute(
                f"""
                SELECT DISTINCT {quoted_column}
                FROM {qualified_table}
                WHERE {quoted_column} IS NOT NULL AND btrim({quoted_column}::text) <> ''
                ORDER BY {quoted_column}
                LIMIT 500
                """
            )
            options[key] = [row[0] for row in cursor.fetchall()]
    return options


def build_portal_where(filters, columns):
    clauses = []
    params = []
    entry_date = quote_name(PORTAL_FILTER_COLUMNS["date_from"])
    exit_date = quote_name(PORTAL_FILTER_COLUMNS["date_to"])
    if filters.get("date_from") and PORTAL_FILTER_COLUMNS["date_from"] in columns:
        clauses.append(f"{entry_date}::date >= %s")
        params.append(filters["date_from"])
    if filters.get("date_to") and PORTAL_FILTER_COLUMNS["date_to"] in columns:
        clauses.append(f"{exit_date}::date <= %s")
        params.append(filters["date_to"])

    for key, column in PORTAL_FILTER_COLUMNS.items():
        if key in {"date_from", "date_to"}:
            continue
        if column in columns and filters.get(key):
            clauses.append(f"{quote_name(column)} = %s")
            params.append(filters[key])
    where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    return where_sql, params


def portal_select_sql(columns):
    select_parts = []
    for column in columns:
        quoted_column = quote_name(column)
        if column in PORTAL_DATE_COLUMNS:
            select_parts.append(f"NULLIF(left({quoted_column}::text, 10), '') AS {quoted_column}")
        else:
            select_parts.append(quoted_column)
    return ", ".join(select_parts)


def portal_column_labels(columns):
    return list(columns)


def portal_rows(params, *, limit=PAGE_SIZE, offset=0):
    columns = get_available_portal_columns()
    filters = {key: params.get(key, "").strip() for key in PORTAL_FILTERS}
    where_sql, query_params = build_portal_where(filters, columns)
    select_columns = portal_select_sql(columns)
    order_sql = f"ORDER BY {quote_name(PORTAL_ORDER_COLUMN)} DESC NULLS LAST"
    qualified_table = f"{quote_name(PORTAL_SCHEMA)}.{quote_name(PORTAL_TABLE)}"

    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {select_columns}
            FROM {qualified_table}
            {where_sql}
            {order_sql}
            LIMIT %s OFFSET %s
            """,
            [*query_params, limit, offset],
        )
        rows = [{"cells": row} for row in cursor.fetchall()]
        cursor.execute(f"SELECT count(*) FROM {qualified_table} {where_sql}", query_params)
        filtered_count = cursor.fetchone()[0]
    return columns, filters, rows, filtered_count


def get_transit_portal_context(params):
    page = clean_page_number(params.get("page"))
    offset = (page - 1) * PAGE_SIZE
    columns, filters, rows, filtered_count = portal_rows(params, limit=PAGE_SIZE, offset=offset)
    total_pages = max(1, (filtered_count + PAGE_SIZE - 1) // PAGE_SIZE)
    if page > total_pages:
        page = total_pages
        offset = (page - 1) * PAGE_SIZE
        columns, filters, rows, filtered_count = portal_rows(params, limit=PAGE_SIZE, offset=offset)
    base_query = QueryDict(mutable=True)
    for key, value in filters.items():
        if value:
            base_query[key] = value
    download_query = base_query.copy()
    download_query["download"] = "xlsx"
    filter_options = portal_filter_options(columns)

    return {
        "columns": columns,
        "column_labels": portal_column_labels(columns),
        "rows": rows,
        "filters": filters,
        "filter_labels": PORTAL_FILTERS,
        "filter_options": filter_options,
        "filter_fields": enrich_filter_fields(PORTAL_FILTER_FIELDS, filters, filter_options),
        "filtered_count": filtered_count,
        "limit": PAGE_SIZE,
        "page": page,
        "total_pages": total_pages,
        "has_previous": page > 1,
        "has_next": page < total_pages,
        "previous_page": page - 1,
        "next_page": page + 1,
        "pagination_items": pagination_items(page, total_pages),
        "page_start": offset + 1 if filtered_count else 0,
        "page_end": min(offset + len(rows), filtered_count),
        "query_string": base_query.urlencode(),
        "download_query_string": download_query.urlencode(),
        "download_limit": DOWNLOAD_LIMIT,
        "eyebrow": "Transit portal database",
        "heading": "Processed Transit Data",
    }


def get_transit_portal_download_rows(params):
    columns = get_available_portal_columns()
    filters = {key: params.get(key, "").strip() for key in PORTAL_FILTERS}
    where_sql, query_params = build_portal_where(filters, columns)
    select_columns = portal_select_sql(columns)
    qualified_table = f"{quote_name(PORTAL_SCHEMA)}.{quote_name(PORTAL_TABLE)}"
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {select_columns}
            FROM {qualified_table}
            {where_sql}
            ORDER BY {quote_name(PORTAL_ORDER_COLUMN)} DESC NULLS LAST
            LIMIT %s
            """,
            [*query_params, DOWNLOAD_LIMIT],
        )
        return portal_column_labels(columns), cursor.fetchall()


def foreign_trucks_qualified_table():
    return f"{quote_name(FOREIGN_TRUCKS_SCHEMA)}.{quote_name(FOREIGN_TRUCKS_MERGED_TABLE)}"


def foreign_trucks_table_exists():
    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT EXISTS (
                SELECT 1
                FROM information_schema.tables
                WHERE table_schema = %s
                  AND table_name = %s
            )
            """,
            [FOREIGN_TRUCKS_SCHEMA, FOREIGN_TRUCKS_MERGED_TABLE],
        )
        return cursor.fetchone()[0]


def get_available_foreign_trucks_columns():
    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = %s
              AND table_name = %s
            """,
            [FOREIGN_TRUCKS_SCHEMA, FOREIGN_TRUCKS_MERGED_TABLE],
        )
        available = {row[0] for row in cursor.fetchall()}
    return [column for column in FOREIGN_TRUCKS_COLUMNS if column in available]


def foreign_trucks_filter_options(columns):
    options = {}
    qualified_table = foreign_trucks_qualified_table()
    with connection.cursor() as cursor:
        for key in [*FOREIGN_TRUCKS_SELECT_FILTERS, *FOREIGN_TRUCKS_COMBOBOX_FILTERS]:
            column = FOREIGN_TRUCKS_FILTER_COLUMNS[key]
            if column not in columns:
                options[key] = []
                continue
            quoted_column = quote_name(column)
            limit = 3000 if key in FOREIGN_TRUCKS_COMBOBOX_FILTERS else 500
            cursor.execute(
                f"""
                SELECT DISTINCT {quoted_column}
                FROM {qualified_table}
                WHERE {quoted_column} IS NOT NULL AND btrim({quoted_column}::text) <> ''
                ORDER BY {quoted_column}
                LIMIT %s
                """,
                [limit],
            )
            options[key] = [row[0] for row in cursor.fetchall()]
    for key in FOREIGN_TRUCKS_TEXT_FILTERS:
        options[key] = []
    return options


def build_foreign_trucks_where(filters, columns):
    clauses = []
    params = []
    date_column = FOREIGN_TRUCKS_FILTER_COLUMNS["date_from"]
    if filters.get("date_from") and date_column in columns:
        clauses.append(f"{quote_name(date_column)}::date >= %s")
        params.append(filters["date_from"])
    if filters.get("date_to") and date_column in columns:
        clauses.append(f"{quote_name(date_column)}::date <= %s")
        params.append(filters["date_to"])

    for key in [*FOREIGN_TRUCKS_SELECT_FILTERS, *FOREIGN_TRUCKS_COMBOBOX_FILTERS]:
        column = FOREIGN_TRUCKS_FILTER_COLUMNS[key]
        if column in columns and filters.get(key):
            clauses.append(f"{quote_name(column)} = %s")
            params.append(filters[key])

    for key in FOREIGN_TRUCKS_TEXT_FILTERS:
        column = FOREIGN_TRUCKS_FILTER_COLUMNS[key]
        if column in columns and filters.get(key):
            clauses.append(f"{quote_name(column)}::text ILIKE %s")
            params.append(f"%{filters[key]}%")

    where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    return where_sql, params


def foreign_trucks_select_sql(columns):
    select_parts = []
    for column in columns:
        quoted_column = quote_name(column)
        if column in FOREIGN_TRUCKS_DATE_COLUMNS:
            select_parts.append(f"NULLIF(left({quoted_column}::text, 19), '') AS {quoted_column}")
        else:
            select_parts.append(quoted_column)
    return ", ".join(select_parts)


def foreign_trucks_column_labels(columns):
    return list(columns)


def foreign_trucks_order_sql(columns):
    order_parts = []
    if "ENTER_DATE" in columns:
        order_parts.append(f"{quote_name('ENTER_DATE')} DESC NULLS LAST")
    if "DATESIGN" in columns:
        order_parts.append(f"{quote_name('DATESIGN')} DESC NULLS LAST")
    if "source_file_path" in columns:
        order_parts.append(f"{quote_name('source_file_path')} DESC")
    return f"ORDER BY {', '.join(order_parts)}" if order_parts else ""


def foreign_trucks_rows(params, *, limit=PAGE_SIZE, offset=0):
    columns = get_available_foreign_trucks_columns()
    filters = {key: params.get(key, "").strip() for key in FOREIGN_TRUCKS_FILTERS}
    if not columns:
        return columns, filters, [], 0
    where_sql, query_params = build_foreign_trucks_where(filters, columns)
    select_columns = foreign_trucks_select_sql(columns)
    order_sql = foreign_trucks_order_sql(columns)
    qualified_table = foreign_trucks_qualified_table()

    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {select_columns}
            FROM {qualified_table}
            {where_sql}
            {order_sql}
            LIMIT %s OFFSET %s
            """,
            [*query_params, limit, offset],
        )
        rows = [{"cells": row} for row in cursor.fetchall()]
        cursor.execute(f"SELECT count(*) FROM {qualified_table} {where_sql}", query_params)
        filtered_count = cursor.fetchone()[0]
    return columns, filters, rows, filtered_count


def get_foreign_trucks_context(params):
    page = clean_page_number(params.get("page"))
    offset = (page - 1) * PAGE_SIZE
    columns, filters, rows, filtered_count = foreign_trucks_rows(params, limit=PAGE_SIZE, offset=offset)
    total_pages = max(1, (filtered_count + PAGE_SIZE - 1) // PAGE_SIZE)
    if page > total_pages:
        page = total_pages
        offset = (page - 1) * PAGE_SIZE
        columns, filters, rows, filtered_count = foreign_trucks_rows(params, limit=PAGE_SIZE, offset=offset)
    base_query = QueryDict(mutable=True)
    for key, value in filters.items():
        if value:
            base_query[key] = value
    download_query = base_query.copy()
    download_query["download"] = "xlsx"
    filter_options = foreign_trucks_filter_options(columns) if columns else {}

    return {
        "columns": columns,
        "column_labels": foreign_trucks_column_labels(columns),
        "rows": rows,
        "filters": filters,
        "filter_labels": FOREIGN_TRUCKS_FILTERS,
        "filter_options": filter_options,
        "filter_fields": enrich_filter_fields(FOREIGN_TRUCKS_FILTER_FIELDS, filters, filter_options),
        "filtered_count": filtered_count,
        "limit": PAGE_SIZE,
        "page": page,
        "total_pages": total_pages,
        "has_previous": page > 1,
        "has_next": page < total_pages,
        "previous_page": page - 1,
        "next_page": page + 1,
        "pagination_items": pagination_items(page, total_pages),
        "page_start": offset + 1 if filtered_count else 0,
        "page_end": min(offset + len(rows), filtered_count),
        "query_string": base_query.urlencode(),
        "download_query_string": download_query.urlencode(),
        "download_limit": DOWNLOAD_LIMIT,
        "eyebrow": "Foreign trucks database",
        "heading": "Raw Foreign Trucks Data",
    }


def get_foreign_trucks_download_rows(params):
    columns = get_available_foreign_trucks_columns()
    filters = {key: params.get(key, "").strip() for key in FOREIGN_TRUCKS_FILTERS}
    if not columns:
        return [], []
    where_sql, query_params = build_foreign_trucks_where(filters, columns)
    select_columns = foreign_trucks_select_sql(columns)
    qualified_table = foreign_trucks_qualified_table()
    order_sql = foreign_trucks_order_sql(columns)
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {select_columns}
            FROM {qualified_table}
            {where_sql}
            {order_sql}
            LIMIT %s
            """,
            [*query_params, DOWNLOAD_LIMIT],
        )
        return foreign_trucks_column_labels(columns), cursor.fetchall()


def dashboard_qualified_table():
    return f"{quote_name(PORTAL_SCHEMA)}.{quote_name(DASHBOARD_TABLE)}"


def dashboard_table_exists():
    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT EXISTS (
                SELECT 1
                FROM information_schema.tables
                WHERE table_schema = %s
                  AND table_name = %s
            )
            """,
            [PORTAL_SCHEMA, DASHBOARD_TABLE],
        )
        return cursor.fetchone()[0]


def getlist_param(params, key):
    if hasattr(params, "getlist"):
        values = params.getlist(key)
    else:
        value = params.get(key, "")
        values = value if isinstance(value, list) else [value]
    return [str(value).strip() for value in values if str(value).strip()]


def dashboard_filters(params):
    filters = {}
    for field in DASHBOARD_FILTER_FIELDS:
        key = field["name"]
        if field["type"] == "multi_select":
            values = getlist_param(params, key)
            filters[key] = values or list(DASHBOARD_DEFAULT_TRANSPORTS)
        else:
            filters[key] = params.get(key, "").strip()
    return filters


def dashboard_filter_options():
    option_fields = [
        field for field in DASHBOARD_FILTER_FIELDS if field["type"] in {"select", "multi_select"}
    ]
    options = {}
    qualified_table = dashboard_qualified_table()
    with connection.cursor() as cursor:
        for field in option_fields:
            key = field["name"]
            column = DASHBOARD_FILTER_COLUMNS[key]
            quoted_column = quote_name(column)
            cursor.execute(
                f"""
                SELECT DISTINCT {quoted_column}
                FROM {qualified_table}
                WHERE {quoted_column} IS NOT NULL AND btrim({quoted_column}::text) <> ''
                ORDER BY {quoted_column}
                LIMIT 500
                """
            )
            values = [row[0] for row in cursor.fetchall()]
            if key == "neqliyyat_novu":
                preferred = [value for value in DASHBOARD_DEFAULT_TRANSPORTS if value in values]
                others = [value for value in values if value not in DASHBOARD_DEFAULT_TRANSPORTS]
                values = preferred + others
            options[key] = values
    return options


def dashboard_filter_fields(filters, options):
    fields = []
    for field in DASHBOARD_FILTER_FIELDS:
        value = filters.get(field["name"], [] if field["type"] == "multi_select" else "")
        fields.append(
            {
                **field,
                "value": value,
                "options": options.get(field["name"], []),
            }
        )
    return fields


def append_in_clause(clauses, params, column, values):
    values = [value for value in values if value]
    if not values:
        return
    placeholders = ", ".join(["%s"] * len(values))
    clauses.append(f"{quote_name(column)} IN ({placeholders})")
    params.extend(values)


def build_dashboard_where(filters, *, include_date=True):
    clauses = []
    params = []
    if include_date:
        if filters.get("date_from"):
            clauses.append(f"{quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s")
            params.append(filters["date_from"])
        if filters.get("date_to"):
            clauses.append(f"{quote_name(DASHBOARD_DATE_COLUMN)}::date <= %s")
            params.append(filters["date_to"])

    append_in_clause(
        clauses,
        params,
        DASHBOARD_FILTER_COLUMNS["neqliyyat_novu"],
        filters.get("neqliyyat_novu", []),
    )
    for key, column in DASHBOARD_FILTER_COLUMNS.items():
        if key in {"date_from", "date_to", "neqliyyat_novu"}:
            continue
        if filters.get(key):
            clauses.append(f"{quote_name(column)} = %s")
            params.append(filters[key])
    where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    return where_sql, params


def fetch_dashboard_totals(filters):
    where_sql, params = build_dashboard_where(filters)
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT
                COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision,
                COUNT(*),
                MIN({quote_name(DASHBOARD_DATE_COLUMN)}::date),
                MAX({quote_name(DASHBOARD_DATE_COLUMN)}::date)
            FROM {dashboard_qualified_table()}
            {where_sql}
            """,
            params,
        )
        total_weight, row_count, min_date, max_date = cursor.fetchone()
    return {
        "total_weight": float(total_weight or 0),
        "row_count": row_count,
        "date_min": min_date,
        "date_max": max_date,
    }


def fetch_grouped_chart(filters, column, *, limit=8, include_date=True):
    where_sql, params = build_dashboard_where(filters, include_date=include_date)
    quoted_column = quote_name(column)
    extra_where = f"{quoted_column} IS NOT NULL AND btrim({quoted_column}::text) <> ''"
    if where_sql:
        where_sql = f"{where_sql} AND {extra_where}"
    else:
        where_sql = f"WHERE {extra_where}"
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {quoted_column}, COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision AS value
            FROM {dashboard_qualified_table()}
            {where_sql}
            GROUP BY {quoted_column}
            ORDER BY value DESC
            LIMIT %s
            """,
            [*params, limit],
        )
        rows = cursor.fetchall()
    return {
        "labels": [row[0] for row in rows],
        "values": [float(row[1] or 0) for row in rows],
    }


def fetch_yearly_trend(filters):
    where_sql, params = build_dashboard_where(filters)
    date_present_clause = f"{quote_name(DASHBOARD_DATE_COLUMN)} IS NOT NULL"
    if where_sql:
        where_sql = f"{where_sql} AND {date_present_clause}"
    else:
        where_sql = f"WHERE {date_present_clause}"
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT
                EXTRACT(YEAR FROM {quote_name(DASHBOARD_DATE_COLUMN)})::int AS year,
                COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision AS value
            FROM {dashboard_qualified_table()}
            {where_sql}
            GROUP BY year
            ORDER BY year
            """,
            params,
        )
        rows = cursor.fetchall()
    return {
        "labels": [str(row[0]) for row in rows],
        "values": [float(row[1] or 0) for row in rows],
    }


def fetch_pair_chart(filters, from_column, to_column, *, limit=8):
    where_sql, params = build_dashboard_where(filters)
    from_expr = f"COALESCE(NULLIF(btrim({quote_name(from_column)}::text), ''), 'Digər')"
    to_expr = f"COALESCE(NULLIF(btrim({quote_name(to_column)}::text), ''), 'Digər')"
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT
                {from_expr} AS source,
                {to_expr} AS target,
                COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision AS value
            FROM {dashboard_qualified_table()}
            {where_sql}
            GROUP BY source, target
            ORDER BY value DESC
            LIMIT %s
            """,
            [*params, limit],
        )
        rows = cursor.fetchall()
    return {
        "labels": [f"{row[0]} -> {row[1]}" for row in rows],
        "sources": [row[0] for row in rows],
        "targets": [row[1] for row in rows],
        "values": [float(row[2] or 0) for row in rows],
    }


def period_total(filters, start_date, end_date):
    where_sql, params = build_dashboard_where(filters, include_date=False)
    date_clause = f"{quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s"
    if where_sql:
        where_sql = f"{where_sql} AND {date_clause}"
    else:
        where_sql = f"WHERE {date_clause}"
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            {where_sql}
            """,
            [*params, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def percent_change(current_value, previous_value):
    if not previous_value:
        return None
    return ((current_value - previous_value) / previous_value) * 100


def format_tons(value):
    return f"{value:,.0f}".replace(",", " ")


def build_comparison_card(title, current_label, previous_label, current_value, previous_value):
    change = percent_change(current_value, previous_value)
    if change is None:
        direction = "neutral"
        change_label = "Müqayisə üçün əvvəlki dövr yoxdur"
    else:
        direction = "up" if change >= 0 else "down"
        sign = "+" if change >= 0 else ""
        change_label = f"{sign}{change:.1f}%"
    return {
        "title": title,
        "current_label": current_label,
        "previous_label": previous_label,
        "current_value": format_tons(current_value),
        "previous_value": format_tons(previous_value),
        "change_label": change_label,
        "direction": direction,
    }


def dashboard_comparison_cards(filters):
    period_metrics = dashboard_period_metrics(filters)
    annual = period_metrics["annual"]
    cards = [
        build_comparison_card(
            "İllik müqayisə",
            annual["label"],
            annual["previous_label"],
            annual["value"],
            annual["previous_value"],
        )
    ]

    monthly = period_metrics.get("monthly")
    if monthly:
        cards.append(
            build_comparison_card(
                "Cari ilin tamamlanmış ayları",
                monthly["label"],
                monthly["previous_label"],
                monthly["value"],
                monthly["previous_value"],
            )
        )
    return cards


def dashboard_period_metrics(filters, periods: TransitDataPeriods | None = None):
    periods = periods or load_transit_data_periods()
    if periods is None:
        return {
            "annual": {
                "title": "Son tamamlanmış il",
                "label": "-",
                "start_date": None,
                "end_date": None,
                "value": 0,
                "value_label": format_tons(0),
                "previous_label": "-",
                "previous_value": 0,
            }
        }

    last_completed_year_total = period_total(
        filters,
        periods.annual_current_start,
        periods.annual_current_end,
    )
    previous_year_total = period_total(
        filters,
        periods.annual_previous_start,
        periods.annual_previous_end,
    )
    current_period_total = period_total(
        filters,
        periods.partial_current_start,
        periods.partial_current_end,
    )
    previous_period_total = period_total(
        filters,
        periods.partial_previous_start,
        periods.partial_previous_end,
    )
    return {
        "annual": {
            "title": "Son tamamlanmış il",
            "label": periods.annual_current_label,
            "start_date": periods.annual_current_start,
            "end_date": periods.annual_current_end,
            "value": last_completed_year_total,
            "value_label": format_tons(last_completed_year_total),
            "previous_label": periods.annual_previous_label,
            "previous_value": previous_year_total,
        },
        "monthly": {
            "title": "Cari tamamlanmış dövr",
            "label": periods.partial_period_label,
            "start_date": periods.partial_current_start,
            "end_date": periods.partial_current_end,
            "value": current_period_total,
            "value_label": format_tons(current_period_total),
            "previous_label": periods.partial_previous_period_label,
            "previous_value": previous_period_total,
        },
    }


def dashboard_period_chart_filters(filters, period_metrics):
    period = period_metrics.get("monthly") or period_metrics["annual"]
    period_filters = {**filters}
    period_filters["date_from"] = period["start_date"].isoformat()
    period_filters["date_to"] = period["end_date"].isoformat()
    return period_filters, period["label"]


def get_transit_dashboard_context(params):
    if not dashboard_table_exists():
        return {
            "available": False,
            "message": (
                "Dashboard table is not available yet. Run "
                "`python manage.py rebuild_transit_dashboard_table` after rebuilding processed transit data."
            ),
        }

    filters = dashboard_filters(params)
    options = dashboard_filter_options()
    totals = fetch_dashboard_totals(filters)
    periods = load_transit_data_periods()
    period_metrics = dashboard_period_metrics(filters, periods)
    period_filters, chart_period_label = dashboard_period_chart_filters(filters, period_metrics)
    corridor_chart = fetch_grouped_chart(period_filters, "Dəhliz", limit=8)
    product_chart = fetch_grouped_chart(period_filters, "Məhsul qrupu (qısaldılmış)", limit=8)
    transport_chart = fetch_grouped_chart(period_filters, "Nəqliyyat növü", limit=8)
    country_flow = fetch_pair_chart(period_filters, "Göndərən ölkə", "Təyinat ölkə", limit=8)
    post_flow = fetch_pair_chart(period_filters, "Gömrük giriş postu", "Gömrük çıxış postu", limit=8)

    return {
        "available": True,
        "filters": filters,
        "filter_fields": dashboard_filter_fields(filters, options),
        "default_transports": DASHBOARD_DEFAULT_TRANSPORTS,
        "summary": {
            **totals,
            "row_count_label": format_tons(totals["row_count"]),
            "top_corridor": corridor_chart["labels"][0] if corridor_chart["labels"] else "Yoxdur",
            "top_product_group": product_chart["labels"][0] if product_chart["labels"] else "Yoxdur",
            "period_metrics": period_metrics,
            "chart_period_label": chart_period_label,
        },
        "comparison_cards": dashboard_comparison_cards(filters),
        "charts": {
            "yearlyTrend": {
                "type": "line",
                "title": "İllik trend",
                **fetch_yearly_trend(filters),
            },
            "transport": {
                "type": "bar",
                "title": "Nəqliyyat növü üzrə",
                **transport_chart,
            },
            "corridor": {
                "type": "bar",
                "title": "Dəhlizlər üzrə",
                **corridor_chart,
            },
            "productGroup": {
                "type": "donut",
                "title": "Məhsul qrupları",
                **product_chart,
            },
            "countryFlow": {
                "type": "flow",
                "title": "Ölkə axını",
                **country_flow,
            },
            "postFlow": {
                "type": "bar",
                "title": "Gömrük postu cütlükləri",
                **post_flow,
            },
        },
    }


def report_period_total(column, value, start_date, end_date):
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            WHERE {quote_name(column)} = %s
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
            """,
            [value, *DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def report_total(start_date, end_date):
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            WHERE {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
            """,
            [*DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def report_total_for_values(column, values, start_date, end_date):
    placeholders = ", ".join(["%s"] * len(values))
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            WHERE {quote_name(column)} IN ({placeholders})
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
            """,
            [*values, *DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def report_top_values(column, start_date, end_date, *, limit=8, preferred_values=None):
    if preferred_values:
        return preferred_values

    quoted_column = quote_name(column)
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {quoted_column}, COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision AS value
            FROM {dashboard_qualified_table()}
            WHERE {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
              AND {quoted_column} IS NOT NULL
              AND btrim({quoted_column}::text) <> ''
            GROUP BY {quoted_column}
            ORDER BY value DESC
            LIMIT %s
            """,
            [*DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date, limit],
        )
        return [row[0] for row in cursor.fetchall()]


def to_thousand_tons(value):
    return value / 1000


def report_percent_change(current_value, previous_value):
    if not previous_value:
        return None
    return ((current_value - previous_value) / previous_value) * 100


def format_report_number(value):
    return f"{to_thousand_tons(value):,.1f}"


def format_report_percent(value):
    if value is None:
        return "-"
    sign = "+" if value >= 0 else ""
    return f"{sign}{value:.1f}%"


def build_report_row(label, values):
    annual_current = values["annual_current"]
    annual_previous = values["annual_previous"]
    partial_current = values["partial_current"]
    partial_previous = values["partial_previous"]
    annual_change = report_percent_change(annual_current, annual_previous)
    partial_change = report_percent_change(partial_current, partial_previous)
    return {
        "label": label,
        "annual_previous": format_report_number(annual_previous),
        "annual_current": format_report_number(annual_current),
        "partial_previous": format_report_number(partial_previous),
        "partial_current": format_report_number(partial_current),
        "annual_change": annual_change,
        "partial_change": partial_change,
        "dynamic": format_report_percent(partial_change),
        "dynamic_direction": "up" if (partial_change or 0) >= 0 else "down",
    }


def report_rows_for_dimension(column, periods: TransitDataPeriods, *, limit=8, preferred_values=None):
    labels = report_top_values(
        column,
        periods.partial_current_start,
        periods.partial_current_end,
        limit=limit,
        preferred_values=preferred_values,
    )
    rows = []
    for label in labels:
        rows.append(
            build_report_row(
                label,
                {
                    "annual_previous": report_period_total(
                        column,
                        label,
                        periods.annual_previous_start,
                        periods.annual_previous_end,
                    ),
                    "annual_current": report_period_total(
                        column,
                        label,
                        periods.annual_current_start,
                        periods.annual_current_end,
                    ),
                    "partial_previous": report_period_total(
                        column,
                        label,
                        periods.partial_previous_start,
                        periods.partial_previous_end,
                    ),
                    "partial_current": report_period_total(
                        column,
                        label,
                        periods.partial_current_start,
                        periods.partial_current_end,
                    ),
                },
            )
        )
    return rows



def report_country_filter_clause():
    return f"({quote_name('Göndərən ölkə')} = %s OR {quote_name('Təyinat ölkə')} = %s)"


def report_period_total_for_country(column, value, start_date, end_date, country):
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            WHERE {quote_name(column)} = %s
              AND {report_country_filter_clause()}
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
            """,
            [value, country, country, *DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def report_total_for_country(start_date, end_date, country):
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            WHERE {report_country_filter_clause()}
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
            """,
            [country, country, *DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def report_total_for_values_and_country(column, values, periods: TransitDataPeriods, country):
    return build_report_row(
        "Total",
        {
            "annual_previous": sum(
                report_period_total_for_country(column, value, periods.annual_previous_start, periods.annual_previous_end, country)
                for value in values
            ),
            "annual_current": sum(
                report_period_total_for_country(column, value, periods.annual_current_start, periods.annual_current_end, country)
                for value in values
            ),
            "partial_previous": sum(
                report_period_total_for_country(column, value, periods.partial_previous_start, periods.partial_previous_end, country)
                for value in values
            ),
            "partial_current": sum(
                report_period_total_for_country(column, value, periods.partial_current_start, periods.partial_current_end, country)
                for value in values
            ),
        },
    )


def report_top_values_for_country(column, start_date, end_date, country, *, limit=8, preferred_values=None):
    if preferred_values:
        return preferred_values

    quoted_column = quote_name(column)
    transport_placeholders = ", ".join(["%s"] * len(DASHBOARD_DEFAULT_TRANSPORTS))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {quoted_column}, COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision AS value
            FROM {dashboard_qualified_table()}
            WHERE {report_country_filter_clause()}
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
              AND {quoted_column} IS NOT NULL
              AND btrim({quoted_column}::text) <> ''
            GROUP BY {quoted_column}
            ORDER BY value DESC
            LIMIT %s
            """,
            [country, country, *DASHBOARD_DEFAULT_TRANSPORTS, start_date, end_date, limit],
        )
        return [row[0] for row in cursor.fetchall()]


def report_rows_for_dimension_and_country(column, periods: TransitDataPeriods, country, *, limit=8, preferred_values=None):
    labels = report_top_values_for_country(
        column,
        periods.partial_current_start,
        periods.partial_current_end,
        country,
        limit=limit,
        preferred_values=preferred_values,
    )
    return [
        build_report_row(
            label,
            {
                "annual_previous": report_period_total_for_country(
                    column, label, periods.annual_previous_start, periods.annual_previous_end, country
                ),
                "annual_current": report_period_total_for_country(
                    column, label, periods.annual_current_start, periods.annual_current_end, country
                ),
                "partial_previous": report_period_total_for_country(
                    column, label, periods.partial_previous_start, periods.partial_previous_end, country
                ),
                "partial_current": report_period_total_for_country(
                    column, label, periods.partial_current_start, periods.partial_current_end, country
                ),
            },
        )
        for label in labels
    ]


TRANSIT_CORRIDOR_REPORT_OPTIONS = ["Şərq-Qərb", "Şimal-Cənub", "Şimal-Qərb", "Cənub-Qərb"]
TRANSIT_CORRIDOR_REPORT_TRANSPORTS = ["Avtomobil", "Dəmiryolu"]


def report_period_total_for_corridor(column, value, start_date, end_date, corridor, *, transport_values=None):
    transport_values = transport_values or DASHBOARD_DEFAULT_TRANSPORTS
    transport_placeholders = ", ".join(["%s"] * len(transport_values))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision
            FROM {dashboard_qualified_table()}
            WHERE {quote_name(column)} = %s
              AND {quote_name('Dəhliz')} = %s
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
            """,
            [value, corridor, *transport_values, start_date, end_date],
        )
        return float(cursor.fetchone()[0] or 0)


def report_total_for_values_and_corridor(column, values, periods: TransitDataPeriods, corridor, *, transport_values=None):
    transport_values = transport_values or DASHBOARD_DEFAULT_TRANSPORTS
    return build_report_row(
        "Total",
        {
            "annual_previous": sum(
                report_period_total_for_corridor(
                    column, value, periods.annual_previous_start, periods.annual_previous_end, corridor, transport_values=transport_values
                )
                for value in values
            ),
            "annual_current": sum(
                report_period_total_for_corridor(
                    column, value, periods.annual_current_start, periods.annual_current_end, corridor, transport_values=transport_values
                )
                for value in values
            ),
            "partial_previous": sum(
                report_period_total_for_corridor(
                    column, value, periods.partial_previous_start, periods.partial_previous_end, corridor, transport_values=transport_values
                )
                for value in values
            ),
            "partial_current": sum(
                report_period_total_for_corridor(
                    column, value, periods.partial_current_start, periods.partial_current_end, corridor, transport_values=transport_values
                )
                for value in values
            ),
        },
    )


def report_top_values_for_corridor(column, start_date, end_date, corridor, *, limit=8, preferred_values=None, transport_values=None):
    if preferred_values:
        return preferred_values

    transport_values = transport_values or DASHBOARD_DEFAULT_TRANSPORTS
    quoted_column = quote_name(column)
    transport_placeholders = ", ".join(["%s"] * len(transport_values))
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT {quoted_column}, COALESCE(SUM({quote_name(DASHBOARD_VALUE_COLUMN)}), 0)::double precision AS value
            FROM {dashboard_qualified_table()}
            WHERE {quote_name('Dəhliz')} = %s
              AND {quote_name('Nəqliyyat növü')} IN ({transport_placeholders})
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date >= %s
              AND {quote_name(DASHBOARD_DATE_COLUMN)}::date < %s
              AND {quoted_column} IS NOT NULL
              AND btrim({quoted_column}::text) <> ''
            GROUP BY {quoted_column}
            ORDER BY value DESC
            LIMIT %s
            """,
            [corridor, *transport_values, start_date, end_date, limit],
        )
        return [row[0] for row in cursor.fetchall()]


def report_rows_for_dimension_and_corridor(
    column, periods: TransitDataPeriods, corridor, *, limit=8, preferred_values=None, transport_values=None
):
    transport_values = transport_values or DASHBOARD_DEFAULT_TRANSPORTS
    labels = report_top_values_for_corridor(
        column,
        periods.partial_current_start,
        periods.partial_current_end,
        corridor,
        limit=limit,
        preferred_values=preferred_values,
        transport_values=transport_values,
    )
    return [
        build_report_row(
            label,
            {
                "annual_previous": report_period_total_for_corridor(
                    column, label, periods.annual_previous_start, periods.annual_previous_end, corridor, transport_values=transport_values
                ),
                "annual_current": report_period_total_for_corridor(
                    column, label, periods.annual_current_start, periods.annual_current_end, corridor, transport_values=transport_values
                ),
                "partial_previous": report_period_total_for_corridor(
                    column, label, periods.partial_previous_start, periods.partial_previous_end, corridor, transport_values=transport_values
                ),
                "partial_current": report_period_total_for_corridor(
                    column, label, periods.partial_current_start, periods.partial_current_end, corridor, transport_values=transport_values
                ),
            },
        )
        for label in labels
    ]


def get_transit_report_country_values():
    with connection.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT country
            FROM (
                SELECT {quote_name('Göndərən ölkə')} AS country FROM {dashboard_qualified_table()}
                UNION
                SELECT {quote_name('Təyinat ölkə')} AS country FROM {dashboard_qualified_table()}
            ) countries
            WHERE country IS NOT NULL AND btrim(country::text) <> ''
            ORDER BY country
            """
        )
        return [row[0] for row in cursor.fetchall()]


def report_summary_totals(periods: TransitDataPeriods):
    values = {
        "annual_previous": report_total(periods.annual_previous_start, periods.annual_previous_end),
        "annual_current": report_total(periods.annual_current_start, periods.annual_current_end),
        "partial_previous": report_total(periods.partial_previous_start, periods.partial_previous_end),
        "partial_current": report_total(periods.partial_current_start, periods.partial_current_end),
    }
    return build_report_row("Total", values)


def report_summary_totals_for_values(column, values, periods: TransitDataPeriods):
    values_map = {
        "annual_previous": report_total_for_values(
            column, values, periods.annual_previous_start, periods.annual_previous_end
        ),
        "annual_current": report_total_for_values(
            column, values, periods.annual_current_start, periods.annual_current_end
        ),
        "partial_previous": report_total_for_values(
            column, values, periods.partial_previous_start, periods.partial_previous_end
        ),
        "partial_current": report_total_for_values(
            column, values, periods.partial_current_start, periods.partial_current_end
        ),
    }
    return build_report_row("Total", values_map)


def transport_narrative_label(label):
    narrative_labels = {
        "Avtomobil": "Avtomobil yolu ilə",
        "Hava": "Hava yolu ilə",
    }
    return narrative_labels.get(label, f"{label} ilə")


def report_sentence(row, period="partial"):
    value_key = "partial_current" if period == "partial" else "annual_current"
    change_key = "partial_change" if period == "partial" else "annual_change"
    change = row[change_key]
    sign = "+" if (change or 0) >= 0 else ""
    label = row["label"]
    return {
        "label": label,
        "narrative_label": transport_narrative_label(label),
        "value": row[value_key],
        "change": "-" if change is None else f"{sign}{change:.1f}%",
        "direction": "up" if (change or 0) >= 0 else "down",
    }


def get_transit_dynamics_report_context(selected_month=None):
    if not dashboard_table_exists():
        return {
            "available": False,
            "message": (
                "Hesabat cədvəli hazır deyil. Əvvəlcə transit_for_dashboard_on_ministry_portal "
                "cədvəlini yenidən yaradın."
            ),
        }

    max_periods = load_transit_data_periods()
    if max_periods is None:
        return {
            "available": False,
            "message": (
                "Hesabat üçün çıxış tarixi tapılmadı. Əvvəlcə transit məlumatlarını idxal edin "
                "və transit_for_dashboard_on_ministry_portal cədvəlini yenidən yaradın."
            ),
        }

    if selected_month is None:
        selected_month = max_periods.reference_month
    selected_month = max(1, min(int(selected_month), max_periods.reference_month))
    periods = compute_transit_data_periods(max_periods.reference_year, selected_month)

    transport_rows = report_rows_for_dimension(
        "Nəqliyyat növü",
        periods,
        preferred_values=DASHBOARD_DEFAULT_TRANSPORTS,
    )
    corridor_rows = report_rows_for_dimension("Dəhliz", periods, limit=REPORT_TOP_LIMIT)
    product_rows = report_rows_for_dimension("Məhsul adı (qısaldılmış)", periods, limit=REPORT_TOP_LIMIT)
    sender_country_rows = report_rows_for_dimension("Göndərən ölkə", periods, limit=REPORT_TOP_LIMIT)
    destination_country_rows = report_rows_for_dimension("Təyinat ölkə", periods, limit=REPORT_TOP_LIMIT)
    total_row = report_summary_totals_for_values("Nəqliyyat növü", DASHBOARD_DEFAULT_TRANSPORTS, periods)

    return {
        "available": True,
        "report_number": format_dynamics_report_number(periods),
        "selected_month": periods.reference_month,
        "max_month": max_periods.reference_month,
        "unit": "min tonla",
        "annual_years": {
            "previous": periods.annual_year_previous,
            "current": periods.annual_year_current,
        },
        "partial_years": {
            "previous": periods.partial_previous_header,
            "current": periods.partial_current_header,
            "label": periods.partial_label,
        },
        "transport": {
            "rows": transport_rows,
            "total": total_row,
            "annual_sentence": report_sentence(total_row, period="annual"),
            "partial_sentence": report_sentence(total_row, period="partial"),
            "bullets_annual": [report_sentence(row, period="annual") for row in transport_rows],
            "bullets_partial": [report_sentence(row, period="partial") for row in transport_rows],
        },
        "corridors": {"rows": corridor_rows},
        "products": {"rows": product_rows},
        "sender_countries": {"rows": sender_country_rows},
        "destination_countries": {"rows": destination_country_rows},
    }


TRANSIT_DYNAMICS_REPORT_TEXTS = {
    "az": {
        "intro": "Tranzit daşımaların əsas dinamika göstəriciləri.",
        "unavailable_title": "Hesabat hazır deyil",
        "unavailable_message": None,
        "report_number_label": "Hesabat nömrəsi",
        "title": "Azərbaycan Üzərindən Tranzit Rejimdə Daşınmış Yüklərin Dinamika Hesabatı",
        "unit": "min tonla",
        "annual_sentence": "{year}-ci ildə tranzit daşımaların həcmi {value} ({change}) min ton təşkil etmişdir.",
        "partial_sentence": "{period} tranzit daşımaların həcmi {value} ({change}) min ton təşkil etmişdir:",
        "value_suffix": "min ton",
        "dynamic": "Dinamika",
        "sections": {
            "transport": {"title": "1. Nəqliyyat növü üzrə tranzit daşımalar", "column": "Nəqliyyat növü"},
            "corridors": {"title": "2. Əsas dəhlizlər üzrə tranzit daşımalar", "column": "Dəhliz"},
            "products": {"title": "3. Əsas məhsullar üzrə tranzit daşımalar", "column": "Məhsul adı (qısaldılmış)"},
            "sender_countries": {"title": "4. Göndərən ölkələr üzrə tranzit daşımalar", "column": "Göndərən ölkə"},
            "destination_countries": {"title": "5. Təyinat ölkələr üzrə tranzit daşımalar", "column": "Təyinat ölkə"},
        },
    },
    "en": {
        "intro": "Main dynamics indicators for transit shipments.",
        "unavailable_title": "Report is not ready",
        "unavailable_message": "The report table is not ready or no reporting date was found. Rebuild the transit dashboard table first.",
        "report_number_label": "Report number",
        "title": "Dynamics Report on Cargo Transported in Transit Mode Through Azerbaijan",
        "unit": "thousand tons",
        "annual_sentence": "In {year}, transit shipment volume amounted to {value} ({change}) thousand tons.",
        "partial_sentence": "In {period}, transit shipment volume amounted to {value} ({change}) thousand tons:",
        "value_suffix": "thousand tons",
        "dynamic": "Dynamics",
        "sections": {
            "transport": {"title": "1. Transit shipments by mode of transport", "column": "Mode of transport"},
            "corridors": {"title": "2. Transit shipments by main corridors", "column": "Corridor"},
            "products": {"title": "3. Transit shipments by main products", "column": "Product name (shortened)"},
            "sender_countries": {"title": "4. Transit shipments by sender countries", "column": "Sender country"},
            "destination_countries": {"title": "5. Transit shipments by destination countries", "column": "Destination country"},
        },
    },
    "ru": {
        "intro": "Основные динамические показатели транзитных перевозок.",
        "unavailable_title": "Отчет не готов",
        "unavailable_message": "Таблица отчета не готова или отчетная дата не найдена. Сначала пересоздайте таблицу транзитного дашборда.",
        "report_number_label": "Номер отчета",
        "title": "Отчет о динамике грузов, перевезенных в транзитном режиме через Азербайджан",
        "unit": "тыс. тонн",
        "annual_sentence": "В {year} году объем транзитных перевозок составил {value} ({change}) тыс. тонн.",
        "partial_sentence": "За {period} объем транзитных перевозок составил {value} ({change}) тыс. тонн:",
        "value_suffix": "тыс. тонн",
        "dynamic": "Динамика",
        "sections": {
            "transport": {"title": "1. Транзитные перевозки по видам транспорта", "column": "Вид транспорта"},
            "corridors": {"title": "2. Транзитные перевозки по основным коридорам", "column": "Коридор"},
            "products": {"title": "3. Транзитные перевозки по основным товарам", "column": "Наименование товара (сокращенное)"},
            "sender_countries": {"title": "4. Транзитные перевозки по странам отправления", "column": "Страна отправления"},
            "destination_countries": {"title": "5. Транзитные перевозки по странам назначения", "column": "Страна назначения"},
        },
    },
}


TRANSIT_TRANSPORT_NARRATIVE_TRANSLATIONS = {
    "az": {
        "Avtomobil": "Avtomobil yolu ilə",
        "Hava": "Hava yolu ilə",
    },
    "en": {
        "Avtomobil": "By road",
        "Boru": "By pipeline",
        "Dəmiryolu": "By railway",
        "Dəniz": "By sea",
        "Hava": "By air",
    },
    "ru": {
        "Avtomobil": "Автомобильным транспортом",
        "Boru": "Трубопроводным транспортом",
        "Dəmiryolu": "Железнодорожным транспортом",
        "Dəniz": "Морским транспортом",
        "Hava": "Воздушным транспортом",
    },
}


def _flatten_short_goods(data):
    return [category["name"] for section in data for category in section["categories"]]


TRANSIT_PRODUCT_TRANSLATIONS = {
    "en": dict(zip(_flatten_short_goods(goods_nomenclature_short_data), _flatten_short_goods(goods_nomenclature_short_english_data))),
    "ru": dict(zip(_flatten_short_goods(goods_nomenclature_short_data), _flatten_short_goods(goods_nomenclature_short_russian_data))),
}


def translate_transit_product_value(value, language):
    if language == "az" or value in (None, ""):
        return value
    return TRANSIT_PRODUCT_TRANSLATIONS.get(language, {}).get(value, value)


def translate_transit_partial_label(label, language):
    if language == "az":
        return label
    parts = str(label).split("-cı ilin ilk ", 1)
    if len(parts) != 2:
        return label
    year, rest = parts
    months = rest.split(" ayında", 1)[0]
    if language == "en":
        return f"the first {months} months of {year}"
    if language == "ru":
        return f"первые {months} месяцев {year} года"
    return label


def translate_transit_report_row(row, dimension, language):
    translated = deepcopy(row)
    if dimension == "Məhsul adı (qısaldılmış)":
        translated["label"] = translate_transit_product_value(row["label"], language)
    elif dimension == "Total":
        translated["label"] = {"az": row["label"], "en": "Total", "ru": "Итого"}.get(language, row["label"])
    else:
        translated["label"] = translate_transit_value(dimension, row["label"], language)
    return translated


def translate_transit_report_sentence(sentence, language):
    translated = deepcopy(sentence)
    label = sentence.get("label")
    translated["label"] = translate_transit_value("Nəqliyyat növü", label, language)
    translated["narrative_label"] = TRANSIT_TRANSPORT_NARRATIVE_TRANSLATIONS.get(language, {}).get(
        label,
        translated["label"],
    )
    return translated


def get_transit_dynamics_report_context_for_language(base_context, language):
    context = deepcopy(base_context)
    text = deepcopy(TRANSIT_DYNAMICS_REPORT_TEXTS[language])
    context["language"] = language
    context["text"] = text

    if not context.get("available"):
        if language != "az":
            context["message"] = text["unavailable_message"]
        return context

    context["unit"] = text["unit"]
    context["partial_years"]["label"] = translate_transit_partial_label(context["partial_years"]["label"], language)
    context["transport"]["annual_summary"] = text["annual_sentence"].format(
        year=context["annual_years"]["current"],
        value=context["transport"]["annual_sentence"]["value"],
        change=context["transport"]["annual_sentence"]["change"],
    )
    context["transport"]["partial_summary"] = text["partial_sentence"].format(
        period=context["partial_years"]["label"],
        value=context["transport"]["partial_sentence"]["value"],
        change=context["transport"]["partial_sentence"]["change"],
    )
    context["transport"]["rows"] = [
        translate_transit_report_row(row, "Nəqliyyat növü", language)
        for row in base_context["transport"]["rows"]
    ]
    context["transport"]["total"] = translate_transit_report_row(base_context["transport"]["total"], "Total", language)
    context["transport"]["annual_sentence"] = translate_transit_report_sentence(
        base_context["transport"]["annual_sentence"], language
    )
    context["transport"]["partial_sentence"] = translate_transit_report_sentence(
        base_context["transport"]["partial_sentence"], language
    )
    context["transport"]["bullets_annual"] = [
        translate_transit_report_sentence(item, language)
        for item in base_context["transport"]["bullets_annual"]
    ]
    context["transport"]["bullets_partial"] = [
        translate_transit_report_sentence(item, language)
        for item in base_context["transport"]["bullets_partial"]
    ]
    corridor_dimension = base_context.get("corridors_dimension", "Dəhliz")
    context["corridors"]["rows"] = [
        translate_transit_report_row(row, corridor_dimension, language)
        for row in base_context["corridors"]["rows"]
    ]
    context["products"]["rows"] = [
        translate_transit_report_row(row, "Məhsul adı (qısaldılmış)", language)
        for row in base_context["products"]["rows"]
    ]
    context["sender_countries"]["rows"] = [
        translate_transit_report_row(row, "Göndərən ölkə", language)
        for row in base_context["sender_countries"]["rows"]
    ]
    context["destination_countries"]["rows"] = [
        translate_transit_report_row(row, "Təyinat ölkə", language)
        for row in base_context["destination_countries"]["rows"]
    ]
    return context


def get_transit_dynamics_report_contexts(base_context):
    return [
        get_transit_dynamics_report_context_for_language(base_context, language)
        for language in ("az", "en", "ru")
    ]



def format_country_report_number(periods: TransitDataPeriods) -> str:
    return f"TR-002-{periods.reference_year}/{periods.reference_month:02d}-Ölkələr"


def get_transit_country_report_context(selected_month=None, selected_country=None):
    if not dashboard_table_exists():
        return {
            "available": False,
            "message": (
                "Hesabat cədvəli hazır deyil. Əvvəlcə transit_for_dashboard_on_ministry_portal "
                "cədvəlini yenidən yaradın."
            ),
        }

    max_periods = load_transit_data_periods()
    if max_periods is None:
        return {
            "available": False,
            "message": (
                "Hesabat üçün çıxış tarixi tapılmadı. Əvvəlcə transit məlumatlarını idxal edin "
                "və transit_for_dashboard_on_ministry_portal cədvəlini yenidən yaradın."
            ),
        }

    countries = get_transit_report_country_values()
    if not countries:
        return {
            "available": False,
            "message": "Ölkələr üzrə hesabat üçün ölkə məlumatı tapılmadı.",
        }

    if selected_country not in countries:
        selected_country = "Türkiyə" if "Türkiyə" in countries else countries[0]

    if selected_month is None:
        selected_month = max_periods.reference_month
    selected_month = max(1, min(int(selected_month), max_periods.reference_month))
    periods = compute_transit_data_periods(max_periods.reference_year, selected_month)

    transport_rows = report_rows_for_dimension_and_country(
        "Nəqliyyat növü",
        periods,
        selected_country,
        preferred_values=DASHBOARD_DEFAULT_TRANSPORTS,
    )
    corridor_rows = report_rows_for_dimension_and_country("Dəhliz", periods, selected_country, limit=REPORT_TOP_LIMIT)
    product_rows = report_rows_for_dimension_and_country(
        "Məhsul adı (qısaldılmış)", periods, selected_country, limit=REPORT_TOP_LIMIT
    )
    sender_country_rows = report_rows_for_dimension_and_country(
        "Göndərən ölkə", periods, selected_country, limit=REPORT_TOP_LIMIT
    )
    destination_country_rows = report_rows_for_dimension_and_country(
        "Təyinat ölkə", periods, selected_country, limit=REPORT_TOP_LIMIT
    )
    total_row = report_total_for_values_and_country("Nəqliyyat növü", DASHBOARD_DEFAULT_TRANSPORTS, periods, selected_country)

    return {
        "available": True,
        "report_type": "countries",
        "report_number": format_country_report_number(periods),
        "selected_month": periods.reference_month,
        "max_month": max_periods.reference_month,
        "selected_country": selected_country,
        "country_options": countries,
        "unit": "min tonla",
        "annual_years": {
            "previous": periods.annual_year_previous,
            "current": periods.annual_year_current,
        },
        "partial_years": {
            "previous": periods.partial_previous_header,
            "current": periods.partial_current_header,
            "label": periods.partial_label,
        },
        "transport": {
            "rows": transport_rows,
            "total": total_row,
            "annual_sentence": report_sentence(total_row, period="annual"),
            "partial_sentence": report_sentence(total_row, period="partial"),
            "bullets_annual": [report_sentence(row, period="annual") for row in transport_rows],
            "bullets_partial": [report_sentence(row, period="partial") for row in transport_rows],
        },
        "corridors": {"rows": corridor_rows},
        "products": {"rows": product_rows},
        "sender_countries": {"rows": sender_country_rows},
        "destination_countries": {"rows": destination_country_rows},
    }


def get_transit_country_report_contexts(base_context):
    reports = get_transit_dynamics_report_contexts(base_context)
    for report in reports:
        if not report.get("available"):
            continue
        language = report["language"]
        country = translate_transit_value("Göndərən ölkə", base_context["selected_country"], language)
        if language == "en":
            report["text"]["title"] = f"Transit Shipments Report Between Azerbaijan and {country}"
        elif language == "ru":
            report["text"]["title"] = f"Отчет о транзитных перевозках: Азербайджан - {country}"
        else:
            report["text"]["title"] = f"Azərbaycan və {country} arasında tranzit daşımalar hesabatı"
    return reports


def format_corridor_report_number(periods: TransitDataPeriods) -> str:
    return f"TR-003-{periods.reference_year}/{periods.reference_month:02d}-Dəhlizlər"


def get_transit_corridor_report_context(selected_month=None, selected_corridor=None):
    if not dashboard_table_exists():
        return {
            "available": False,
            "message": (
                "Hesabat cədvəli hazır deyil. Əvvəlcə transit_for_dashboard_on_ministry_portal "
                "cədvəlini yenidən yaradın."
            ),
        }

    max_periods = load_transit_data_periods()
    if max_periods is None:
        return {
            "available": False,
            "message": (
                "Hesabat üçün çıxış tarixi tapılmadı. Əvvəlcə transit məlumatlarını idxal edin "
                "və transit_for_dashboard_on_ministry_portal cədvəlini yenidən yaradın."
            ),
        }

    if selected_corridor not in TRANSIT_CORRIDOR_REPORT_OPTIONS:
        selected_corridor = TRANSIT_CORRIDOR_REPORT_OPTIONS[0]

    if selected_month is None:
        selected_month = max_periods.reference_month
    selected_month = max(1, min(int(selected_month), max_periods.reference_month))
    periods = compute_transit_data_periods(max_periods.reference_year, selected_month)

    transport_rows = report_rows_for_dimension_and_corridor(
        "Nəqliyyat növü",
        periods,
        selected_corridor,
        preferred_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS,
        transport_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS,
    )
    direction_rows = report_rows_for_dimension_and_corridor(
        "Dəhliz (istiqamətlə)", periods, selected_corridor, limit=REPORT_TOP_LIMIT, transport_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS
    )
    product_rows = report_rows_for_dimension_and_corridor(
        "Məhsul adı (qısaldılmış)", periods, selected_corridor, limit=REPORT_TOP_LIMIT, transport_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS
    )
    sender_country_rows = report_rows_for_dimension_and_corridor(
        "Göndərən ölkə", periods, selected_corridor, limit=REPORT_TOP_LIMIT, transport_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS
    )
    destination_country_rows = report_rows_for_dimension_and_corridor(
        "Təyinat ölkə", periods, selected_corridor, limit=REPORT_TOP_LIMIT, transport_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS
    )
    total_row = report_total_for_values_and_corridor(
        "Nəqliyyat növü", TRANSIT_CORRIDOR_REPORT_TRANSPORTS, periods, selected_corridor,
        transport_values=TRANSIT_CORRIDOR_REPORT_TRANSPORTS,
    )

    return {
        "available": True,
        "report_type": "corridors",
        "report_number": format_corridor_report_number(periods),
        "selected_month": periods.reference_month,
        "max_month": max_periods.reference_month,
        "selected_corridor": selected_corridor,
        "corridor_options": TRANSIT_CORRIDOR_REPORT_OPTIONS,
        "corridors_dimension": "Dəhliz (istiqamətlə)",
        "unit": "min tonla",
        "annual_years": {
            "previous": periods.annual_year_previous,
            "current": periods.annual_year_current,
        },
        "partial_years": {
            "previous": periods.partial_previous_header,
            "current": periods.partial_current_header,
            "label": periods.partial_label,
        },
        "transport": {
            "rows": transport_rows,
            "total": total_row,
            "annual_sentence": report_sentence(total_row, period="annual"),
            "partial_sentence": report_sentence(total_row, period="partial"),
            "bullets_annual": [report_sentence(row, period="annual") for row in transport_rows],
            "bullets_partial": [report_sentence(row, period="partial") for row in transport_rows],
        },
        "corridors": {"rows": direction_rows},
        "products": {"rows": product_rows},
        "sender_countries": {"rows": sender_country_rows},
        "destination_countries": {"rows": destination_country_rows},
    }


def get_transit_corridor_report_contexts(base_context):
    reports = get_transit_dynamics_report_contexts(base_context)
    for report in reports:
        if not report.get("available"):
            continue
        language = report["language"]
        corridor = translate_transit_value("Dəhliz", base_context["selected_corridor"], language)
        if language == "en":
            report["text"]["title"] = f"Transit Shipments Report for the {corridor} Corridor"
            report["text"]["sections"]["corridors"] = {
                "title": "2. Transit shipments by direction within the corridor",
                "column": "Corridor direction",
            }
        elif language == "ru":
            report["text"]["title"] = f"Отчет о транзитных перевозках по коридору {corridor}"
            report["text"]["sections"]["corridors"] = {
                "title": "2. Транзитные перевозки по направлениям в коридоре",
                "column": "Направление коридора",
            }
        else:
            report["text"]["title"] = f"{corridor} dəhlizi üzrə tranzit daşımalar hesabatı"
            report["text"]["sections"]["corridors"] = {
                "title": "2. Dəhlizdə istiqamət üzrə tranzit daşımalar",
                "column": "Dəhliz (istiqamətlə)",
            }
    return reports


TRANSIT_POST_REPORT_TEXTS = {
    "az": {
        "intro": "Tranzit daşımaların postlar üzrə əsas göstəriciləri.",
        "report_number_label": "Hesabat nömrəsi",
        "title": "Azərbaycan üzərindən tranzit rejimdə daşınmış yüklərin postlar üzrə hesabatı",
        "unit": "min tonla",
        "dynamic": "Dinamika",
        "sections": {
            "entry_posts": {"title": "1. Gömrük giriş postları üzrə tranzit daşımalar", "column": "Gömrük giriş postu"},
            "exit_posts": {"title": "2. Gömrük çıxış postları üzrə tranzit daşımalar", "column": "Gömrük çıxış postu"},
        },
    },
    "en": {
        "intro": "Main indicators for transit shipments by customs posts.",
        "report_number_label": "Report number",
        "title": "Report on Cargo Transported in Transit Mode Through Azerbaijan by Customs Posts",
        "unit": "thousand tons",
        "dynamic": "Dynamics",
        "sections": {
            "entry_posts": {"title": "1. Transit shipments by customs entry posts", "column": "Customs entry post"},
            "exit_posts": {"title": "2. Transit shipments by customs exit posts", "column": "Customs exit post"},
        },
    },
    "ru": {
        "intro": "Основные показатели транзитных перевозок по таможенным постам.",
        "report_number_label": "Номер отчета",
        "title": "Отчет о грузах, перевезенных в транзитном режиме через Азербайджан, по таможенным постам",
        "unit": "тыс. тонн",
        "dynamic": "Динамика",
        "sections": {
            "entry_posts": {"title": "1. Транзитные перевозки по таможенным постам въезда", "column": "Таможенный пост въезда"},
            "exit_posts": {"title": "2. Транзитные перевозки по таможенным постам выезда", "column": "Таможенный пост выезда"},
        },
    },
}


def format_posts_report_number(periods: TransitDataPeriods) -> str:
    return f"TR-004-{periods.reference_year}/{periods.reference_month:02d}-Postlar"


def get_transit_posts_report_context(selected_month=None):
    if not dashboard_table_exists():
        return {
            "available": False,
            "message": (
                "Hesabat cədvəli hazır deyil. Əvvəlcə transit_for_dashboard_on_ministry_portal "
                "cədvəlini yenidən yaradın."
            ),
        }

    max_periods = load_transit_data_periods()
    if max_periods is None:
        return {
            "available": False,
            "message": (
                "Hesabat üçün çıxış tarixi tapılmadı. Əvvəlcə transit məlumatlarını idxal edin "
                "və transit_for_dashboard_on_ministry_portal cədvəlini yenidən yaradın."
            ),
        }

    if selected_month is None:
        selected_month = max_periods.reference_month
    selected_month = max(1, min(int(selected_month), max_periods.reference_month))
    periods = compute_transit_data_periods(max_periods.reference_year, selected_month)

    return {
        "available": True,
        "report_type": "posts",
        "report_number": format_posts_report_number(periods),
        "selected_month": periods.reference_month,
        "max_month": max_periods.reference_month,
        "unit": "min tonla",
        "annual_years": {
            "previous": periods.annual_year_previous,
            "current": periods.annual_year_current,
        },
        "partial_years": {
            "previous": periods.partial_previous_header,
            "current": periods.partial_current_header,
            "label": periods.partial_label,
        },
        "entry_posts": {
            "rows": report_rows_for_dimension("Gömrük giriş postu", periods, limit=REPORT_POST_LIMIT),
        },
        "exit_posts": {
            "rows": report_rows_for_dimension("Gömrük çıxış postu", periods, limit=REPORT_POST_LIMIT),
        },
    }


def get_transit_posts_report_context_for_language(base_context, language):
    context = deepcopy(base_context)
    text = deepcopy(TRANSIT_POST_REPORT_TEXTS[language])
    context["language"] = language
    context["text"] = text

    if not context.get("available"):
        if language != "az":
            context["message"] = TRANSIT_DYNAMICS_REPORT_TEXTS[language]["unavailable_message"]
        return context

    context["unit"] = text["unit"]
    context["partial_years"]["label"] = translate_transit_partial_label(context["partial_years"]["label"], language)
    context["entry_posts"]["rows"] = [
        translate_transit_report_row(row, "Gömrük giriş postu", language)
        for row in base_context["entry_posts"]["rows"]
    ]
    context["exit_posts"]["rows"] = [
        translate_transit_report_row(row, "Gömrük çıxış postu", language)
        for row in base_context["exit_posts"]["rows"]
    ]
    return context


def get_transit_posts_report_contexts(base_context):
    return [
        get_transit_posts_report_context_for_language(base_context, language)
        for language in ("az", "en", "ru")
    ]


DOCX_HEADER_FILL = "073763"
DOCX_POSITIVE_PERCENT_COLOR = "00B050"
DOCX_NEGATIVE_PERCENT_COLOR = "FF0000"


def _set_docx_run_style(run, *, bold=False, size=12, color="073763"):
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _dynamic_percent_color(direction):
    return DOCX_NEGATIVE_PERCENT_COLOR if direction == "down" else DOCX_POSITIVE_PERCENT_COLOR


def _append_docx_parenthesized_change(paragraph, change, direction, *, bold=False, size=12):
    if change == "-":
        run = paragraph.add_run(f" ({change})")
        _set_docx_run_style(run, bold=bold, size=size)
        return
    open_run = paragraph.add_run(" (")
    _set_docx_run_style(open_run, bold=bold, size=size)
    change_run = paragraph.add_run(change)
    _set_docx_run_style(change_run, bold=True, size=size, color=_dynamic_percent_color(direction))
    close_run = paragraph.add_run(")")
    _set_docx_run_style(close_run, bold=bold, size=size)


def _add_docx_paragraph_with_change(
    document,
    before_change,
    change,
    direction,
    after_change="",
    *,
    bold=False,
    size=12,
    space_after=6,
    list_style=None,
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    if list_style is not None:
        paragraph.style = list_style
    run = paragraph.add_run(before_change)
    _set_docx_run_style(run, bold=bold, size=size)
    _append_docx_parenthesized_change(paragraph, change, direction, bold=bold, size=size)
    if after_change:
        after_run = paragraph.add_run(after_change)
        _set_docx_run_style(after_run, bold=bold, size=size)
    return paragraph


def _add_docx_paragraph(document, text, *, bold=False, size=12, alignment=None, space_after=6):
    paragraph = document.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    _set_docx_run_style(run, bold=bold, size=size)
    return paragraph


def _shade_docx_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))


def _style_docx_cell(cell, *, bold=False, color="073763", fill=None, align_center=False):
    if fill:
        _shade_docx_cell(cell, fill)
    if align_center:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        if align_center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            _set_docx_run_style(run, bold=bold, size=12, color=color)


def _style_docx_table_value_cell(cell, value, *, row_data=None, cell_index=0, bold=False, fill=None):
    cell.text = str(value)
    is_dynamic_column = cell_index == 5
    if is_dynamic_column and value != "-" and row_data is not None:
        _style_docx_cell(
            cell,
            bold=True,
            color=_dynamic_percent_color(row_data["dynamic_direction"]),
            fill=fill,
            align_center=True,
        )
        return
    _style_docx_cell(
        cell,
        bold=bold,
        fill=fill,
        align_center=cell_index > 0,
    )


def _set_report_docx_table_widths(table):
    column_widths = [Inches(2)] + [Inches(1.1)] * (len(table.columns) - 1)
    table.autofit = False
    for column_index, width in enumerate(column_widths):
        table.columns[column_index].width = width
        for cell in table.columns[column_index].cells:
            cell.width = width


def _add_report_docx_table(document, headers, rows, total_row=None):
    table_rows = len(rows) + 1 + (1 if total_row else 0)
    table = document.add_table(rows=table_rows, cols=len(headers))
    table.style = "Table Grid"
    _set_report_docx_table_widths(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        _style_docx_cell(
            cell,
            bold=True,
            color="FFFFFF",
            fill=DOCX_HEADER_FILL,
            align_center=True,
        )

    for row_index, row in enumerate(rows, start=1):
        values = [
            row["label"],
            row["annual_previous"],
            row["annual_current"],
            row["partial_previous"],
            row["partial_current"],
            row["dynamic"],
        ]
        for cell_index, value in enumerate(values):
            cell = table.rows[row_index].cells[cell_index]
            _style_docx_table_value_cell(cell, value, row_data=row, cell_index=cell_index)

    if total_row:
        values = [
            total_row["label"],
            total_row["annual_previous"],
            total_row["annual_current"],
            total_row["partial_previous"],
            total_row["partial_current"],
            total_row["dynamic"],
        ]
        for cell_index, value in enumerate(values):
            cell = table.rows[-1].cells[cell_index]
            _style_docx_table_value_cell(
                cell,
                value,
                row_data=total_row,
                cell_index=cell_index,
                bold=True,
                fill="EAF2FB",
            )

    document.add_paragraph()


def _add_transit_docx_summary(document, report_context, *, period):
    language = report_context.get("language", "az")
    transport = report_context["transport"]
    text = report_context.get("text", TRANSIT_DYNAMICS_REPORT_TEXTS["az"])
    value_suffix = text["value_suffix"]

    if period == "annual":
        sentence = transport["annual_sentence"]
        if language == "en":
            before_change = (
                f"In {report_context['annual_years']['current']}, transit shipment volume "
                f"amounted to {sentence['value']} "
            )
            after_change = f" {value_suffix}."
        elif language == "ru":
            before_change = (
                f"В {report_context['annual_years']['current']} году объем транзитных перевозок "
                f"составил {sentence['value']} "
            )
            after_change = f" {value_suffix}."
        else:
            before_change = (
                f"{report_context['annual_years']['current']}-ci ildə tranzit daşımaların həcmi "
                f"{sentence['value']} "
            )
            after_change = " min ton təşkil etmişdir."
    else:
        sentence = transport["partial_sentence"]
        if language == "en":
            before_change = (
                f"In {report_context['partial_years']['label']}, transit shipment volume "
                f"amounted to {sentence['value']} "
            )
            after_change = f" {value_suffix}:"
        elif language == "ru":
            before_change = (
                f"За {report_context['partial_years']['label']} объем транзитных перевозок "
                f"составил {sentence['value']} "
            )
            after_change = f" {value_suffix}:"
        else:
            before_change = (
                f"{report_context['partial_years']['label']} tranzit daşımaların həcmi "
                f"{sentence['value']} "
            )
            after_change = " min ton təşkil etmişdir:"

    _add_docx_paragraph_with_change(
        document,
        before_change,
        sentence["change"],
        sentence["direction"],
        after_change,
        bold=True,
    )


def _add_posts_report_docx_table(document, report_context, section_key, rows):
    text = report_context["text"]
    section_text = text["sections"][section_key]
    period_headers = [
        report_context["annual_years"]["previous"],
        report_context["annual_years"]["current"],
        report_context["partial_years"]["previous"],
        report_context["partial_years"]["current"],
    ]
    _add_docx_paragraph(
        document,
        f"{section_text['title']}, {report_context['unit']}",
        bold=True,
        size=14,
        space_after=10,
    )
    _add_report_docx_table(document, [section_text["column"], *period_headers, text["dynamic"]], rows)


def build_transit_posts_report_docx(report_context):
    text = report_context.get("text", TRANSIT_POST_REPORT_TEXTS["az"])
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    _add_docx_paragraph(
        document,
        f"{text['report_number_label']}: {report_context['report_number']}",
        bold=True,
        size=10,
        space_after=18,
    )
    _add_docx_paragraph(
        document,
        text["title"],
        bold=True,
        size=16,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=28,
    )
    _add_posts_report_docx_table(document, report_context, "entry_posts", report_context["entry_posts"]["rows"])
    _add_posts_report_docx_table(document, report_context, "exit_posts", report_context["exit_posts"]["rows"])

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def build_transit_dynamics_report_docx(report_context):
    if report_context.get("report_type") == "posts":
        return build_transit_posts_report_docx(report_context)

    text = report_context.get("text", TRANSIT_DYNAMICS_REPORT_TEXTS["az"])
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    _add_docx_paragraph(
        document,
        f"{text['report_number_label']}: {report_context['report_number']}",
        bold=True,
        size=10,
        space_after=18,
    )
    _add_docx_paragraph(
        document,
        text["title"],
        bold=True,
        size=16,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=28,
    )

    transport = report_context["transport"]
    _add_docx_paragraph(
        document,
        f"{text['sections']['transport']['title']}, {report_context['unit']}",
        bold=True,
        size=14,
        space_after=10,
    )
    _add_transit_docx_summary(document, report_context, period="annual")
    for item in transport["bullets_annual"]:
        _add_docx_paragraph_with_change(
            document,
            f"{item.get('narrative_label', item['label'])}: {item['value']} {text['value_suffix']} ",
            item["change"],
            item["direction"],
            space_after=2,
            list_style=document.styles["List Bullet"],
        )
    _add_transit_docx_summary(document, report_context, period="partial")
    for item in transport["bullets_partial"]:
        _add_docx_paragraph_with_change(
            document,
            f"{item.get('narrative_label', item['label'])}: {item['value']} {text['value_suffix']} ",
            item["change"],
            item["direction"],
            space_after=2,
            list_style=document.styles["List Bullet"],
        )

    period_headers = [
        report_context["annual_years"]["previous"],
        report_context["annual_years"]["current"],
        report_context["partial_years"]["previous"],
        report_context["partial_years"]["current"],
    ]
    _add_report_docx_table(
        document,
        [text["sections"]["transport"]["column"], *period_headers, text["dynamic"]],
        transport["rows"],
        total_row=transport["total"],
    )

    report_sections = [
        (text["sections"]["corridors"], report_context["corridors"]["rows"]),
        (text["sections"]["products"], report_context["products"]["rows"]),
        (text["sections"]["sender_countries"], report_context["sender_countries"]["rows"]),
        (text["sections"]["destination_countries"], report_context["destination_countries"]["rows"]),
    ]
    for section_text, rows in report_sections:
        _add_docx_paragraph(
            document,
            f"{section_text['title']}, {report_context['unit']}",
            bold=True,
            size=14,
            space_after=10,
        )
        _add_report_docx_table(document, [section_text["column"], *period_headers, text["dynamic"]], rows)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def build_transit_dynamics_report_pdf(report_context):
    import dxpdf

    docx_output = build_transit_dynamics_report_docx(report_context)
    pdf_bytes = dxpdf.convert(docx_output.getvalue())
    output = BytesIO(pdf_bytes)
    output.seek(0)
    return output
