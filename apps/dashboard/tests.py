from datetime import date
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.urls import reverse
from docx import Document

from apps.catalog.models import Module, Submodule
from apps.dashboard.services import (
    FOREIGN_TRUCKS_FILTER_COLUMNS,
    FOREIGN_TRUCKS_FILTER_FIELDS,
    LOCAL_TRUCKS_FILTER_COLUMNS,
    LOCAL_TRUCKS_FILTER_FIELDS,
    available_filter_fields,
    build_transit_dynamics_report_docx,
    build_transit_dynamics_report_pdf,
    format_report_percent,
    transport_narrative_label,
)
from apps.dashboard.transit_periods import compute_transit_data_periods, format_dynamics_report_number


class TruckRawDataTests(TestCase):
    def test_foreign_truck_filter_fields_keep_all_existing_columns(self):
        columns = {"ENTER_DATE", "IDN", "FROMTO", "SHORT_NAME", "PERM_BLANK_NO", "HES_NAME", "CONS_NAME"}

        fields = available_filter_fields(
            FOREIGN_TRUCKS_FILTER_FIELDS,
            FOREIGN_TRUCKS_FILTER_COLUMNS,
            columns,
        )

        field_names = [field["name"] for field in fields]
        self.assertEqual(
            field_names,
            ["date_from", "date_to", "idn", "perm_blank_no", "fromto", "short_name", "hes_name", "cons_name"],
        )

    def test_local_truck_filter_fields_hide_requested_fields(self):
        columns = {"ENTER_DATE", "DATESIGN", "IDN", "FROMTO", "SHORT_NAME", "PERM_BLANK_NO", "HES_NAME", "CONS_NAME"}

        fields = available_filter_fields(
            LOCAL_TRUCKS_FILTER_FIELDS,
            LOCAL_TRUCKS_FILTER_COLUMNS,
            columns,
        )

        field_names = [field["name"] for field in fields]
        self.assertEqual(field_names, ["date_from", "date_to", "idn", "fromto", "datesign_from", "datesign_to"])
        self.assertNotIn("short_name", field_names)
        self.assertNotIn("perm_blank_no", field_names)
        self.assertNotIn("hes_name", field_names)
        self.assertNotIn("cons_name", field_names)

    @patch("apps.dashboard.views.get_local_trucks_context")
    def test_yerli_tirlar_submodule_uses_local_trucks_context(self, get_local_trucks_context):
        module = Module.objects.create(name="Raw data", slug="datalar")
        submodule = Submodule.objects.create(module=module, name="Yerli TIR-lar", slug="yerli-tir-lar")
        user = get_user_model().objects.create_user(username="dashboard-user", password="pass")
        self.client.force_login(user)
        get_local_trucks_context.return_value = {
            "columns": [],
            "column_labels": [],
            "rows": [],
            "filters": {},
            "filter_labels": {},
            "filter_options": {},
            "filter_fields": [],
            "filtered_count": 0,
            "limit": 50,
            "page": 1,
            "total_pages": 1,
            "has_previous": False,
            "has_next": False,
            "previous_page": 0,
            "next_page": 2,
            "pagination_items": [{"number": 1, "current": True}],
            "page_start": 0,
            "page_end": 0,
            "query_string": "",
            "download_query_string": "download=xlsx",
            "download_limit": 10000,
            "eyebrow": "Local trucks database",
            "heading": "Raw Local Trucks Data",
        }

        response = self.client.get(
            reverse(
                "dashboard:submodule_detail",
                kwargs={"module_slug": module.slug, "submodule_slug": submodule.slug},
            )
        )

        self.assertEqual(response.status_code, 200)
        get_local_trucks_context.assert_called_once()
        self.assertContains(response, "Raw Local Trucks Data")


class TransitDataPeriodsTests(TestCase):
    def test_compute_periods_from_latest_data_may_2026(self):
        periods = compute_transit_data_periods(2026, 5)

        self.assertEqual(periods.reference_year, 2026)
        self.assertEqual(periods.reference_month, 5)
        self.assertEqual(periods.annual_year_current, 2025)
        self.assertEqual(periods.annual_year_previous, 2024)
        self.assertEqual(periods.partial_year_current, 2026)
        self.assertEqual(periods.partial_year_previous, 2025)

        self.assertEqual(periods.annual_current_start, date(2025, 1, 1))
        self.assertEqual(periods.annual_current_end, date(2026, 1, 1))
        self.assertEqual(periods.annual_previous_start, date(2024, 1, 1))
        self.assertEqual(periods.annual_previous_end, date(2025, 1, 1))

        self.assertEqual(periods.partial_current_start, date(2026, 1, 1))
        self.assertEqual(periods.partial_current_end, date(2026, 6, 1))
        self.assertEqual(periods.partial_previous_start, date(2025, 1, 1))
        self.assertEqual(periods.partial_previous_end, date(2025, 6, 1))

        self.assertEqual(periods.partial_label, "2026-cı ilin ilk 5 ayında")
        self.assertEqual(periods.partial_period_label, "2026 Yan-05")

    def test_compute_periods_handles_december(self):
        periods = compute_transit_data_periods(2025, 12)

        self.assertEqual(periods.partial_current_end, date(2026, 1, 1))
        self.assertEqual(periods.partial_previous_end, date(2025, 1, 1))

    def test_format_dynamics_report_number(self):
        periods = compute_transit_data_periods(2026, 5)
        self.assertEqual(format_dynamics_report_number(periods), "TR-001-2026/05-Dinamika")


class TransportNarrativeLabelTests(TestCase):
    def test_transport_narrative_labels(self):
        self.assertEqual(transport_narrative_label("Avtomobil"), "Avtomobil yolu ilə")
        self.assertEqual(transport_narrative_label("Hava"), "Hava yolu ilə")
        self.assertEqual(transport_narrative_label("Dəmiryolu"), "Dəmiryolu ilə")


class ReportPercentFormattingTests(TestCase):
    def test_format_report_percent_uses_one_decimal(self):
        self.assertEqual(format_report_percent(8.44), "+8.4%")
        self.assertEqual(format_report_percent(-3.16), "-3.2%")
        self.assertEqual(format_report_percent(None), "-")


class TransitDynamicsDocxTests(TestCase):
    def test_builds_word_report_from_context(self):
        row = {
            "label": "Dəmir yolu",
            "annual_previous": "10.0",
            "annual_current": "12.0",
            "partial_previous": "4.0",
            "partial_current": "5.0",
            "dynamic": "+25.0%",
            "dynamic_direction": "up",
        }
        context = {
            "available": True,
            "report_number": "TR-001-2026/05-Dinamika",
            "unit": "min tonla",
            "annual_years": {"previous": 2024, "current": 2025},
            "partial_years": {"previous": "2025*", "current": "2026*", "label": "2026-cı ilin ilk 5 ayında"},
            "transport": {
                "rows": [row],
                "total": {**row, "label": "Total"},
                "annual_sentence": {"value": "12.0", "change": "+20.0%", "direction": "up"},
                "partial_sentence": {"value": "5.0", "change": "+25.0%", "direction": "up"},
                "bullets_annual": [
                    {"label": "Dəmir yolu", "value": "12.0", "change": "+20.0%", "direction": "up"}
                ],
                "bullets_partial": [
                    {"label": "Dəmir yolu", "value": "5.0", "change": "+25.0%", "direction": "up"}
                ],
            },
            "corridors": {"rows": [row]},
            "products": {"rows": [row]},
            "sender_countries": {"rows": [row]},
            "destination_countries": {"rows": [row]},
        }

        output = build_transit_dynamics_report_docx(context)
        document = Document(output)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("Hesabat nömrəsi: TR-001-2026/05-Dinamika", text)
        self.assertIn("Azərbaycan Üzərindən Tranzit Rejimdə Daşınmış Yüklərin Dinamika Hesabatı", text)
        self.assertEqual(len(document.tables), 5)

    def test_builds_pdf_report_from_word(self):
        row = {
            "label": "Dəmir yolu",
            "annual_previous": "10.0",
            "annual_current": "12.0",
            "partial_previous": "4.0",
            "partial_current": "5.0",
            "dynamic": "+25.0%",
            "dynamic_direction": "up",
        }
        context = {
            "available": True,
            "report_number": "TR-001-2026/05-Dinamika",
            "unit": "min tonla",
            "annual_years": {"previous": 2024, "current": 2025},
            "partial_years": {"previous": "2025*", "current": "2026*", "label": "2026-cı ilin ilk 5 ayında"},
            "transport": {
                "rows": [row],
                "total": {**row, "label": "Total"},
                "annual_sentence": {"value": "12.0", "change": "+20.0%", "direction": "up"},
                "partial_sentence": {"value": "5.0", "change": "+25.0%", "direction": "up"},
                "bullets_annual": [
                    {"label": "Dəmir yolu", "value": "12.0", "change": "+20.0%", "direction": "up"}
                ],
                "bullets_partial": [
                    {"label": "Dəmir yolu", "value": "5.0", "change": "+25.0%", "direction": "up"}
                ],
            },
            "corridors": {"rows": [row]},
            "products": {"rows": [row]},
            "sender_countries": {"rows": [row]},
            "destination_countries": {"rows": [row]},
        }

        output = build_transit_dynamics_report_pdf(context)
        self.assertTrue(output.getvalue().startswith(b"%PDF"))
